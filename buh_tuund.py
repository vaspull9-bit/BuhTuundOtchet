#======================================================================
# BuhTuundOtchet v7.3.0 - Доработка сохранения БД
import sys
import os
import sqlite3
import re
import shutil
import io
import calendar
from datetime import datetime
from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from PyQt6.QtGui import *
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment
from openpyxl.drawing.image import Image as ExcelImage

# ==================== БАЗА ДАННЫХ ====================
class DatabaseManager:
    def __init__(self, db_path='buh_tuund.db'):
        self.conn = sqlite3.connect(db_path, check_same_thread=False)
        self.create_tables()
        
    def create_tables(self):
        cursor = self.conn.cursor()

        # Основная таблица reports
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS reports (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company TEXT,
                period_start TEXT,
                period_end TEXT,
                doc_type TEXT,
                product_group TEXT,
                nomenclature TEXT,
                revenue REAL,
                cost_price REAL,
                gross_profit REAL,
                sales_expenses REAL,
                other_income_expenses REAL,
                net_profit REAL,
                vat_deductible REAL,
                vat_to_budget REAL,
                quantity INTEGER,
                import_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        # Добавляем новые колонки, если их нет
        cursor.execute("PRAGMA table_info(reports)")
        existing = [col[1] for col in cursor.fetchall()]
        
        new_columns = {
            'seller': 'TEXT',
            'buyer': 'TEXT',
            'document_number': 'TEXT',
            'document_date': 'TEXT',
            'operation_code': 'TEXT',
            'acceptance_date': 'TEXT',
            'payment_document': 'TEXT',
            'purchase_amount_with_vat': 'REAL',
            'sales_amount_without_vat': 'REAL',
            'sales_amount_with_vat': 'REAL'
        }
        
        for col, typ in new_columns.items():
            if col not in existing:
                cursor.execute(f"ALTER TABLE reports ADD COLUMN {col} {typ}")

        cursor.execute('''
            CREATE TABLE IF NOT EXISTS import_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT,
                records_count INTEGER,
                import_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        self.conn.commit()

    def save_data(self, df):
        """Сохраняет данные из DataFrame в таблицу reports"""
        df_to_save = df.copy()

        # Все возможные колонки со значениями по умолчанию
        all_columns = {
            'company': '',
            'period_start': '',
            'period_end': '',
            'doc_type': '',
            'product_group': '',
            'nomenclature': '',
            'revenue': 0.0,
            'cost_price': 0.0,
            'gross_profit': 0.0,
            'sales_expenses': 0.0,
            'other_income_expenses': 0.0,
            'net_profit': 0.0,
            'vat_deductible': 0.0,
            'vat_to_budget': 0.0,
            'quantity': 0,
            'seller': '',
            'buyer': '',
            'document_number': '',
            'document_date': '',
            'operation_code': '',
            'acceptance_date': '',
            'payment_document': '',
            'purchase_amount_with_vat': 0.0,
            'sales_amount_without_vat': 0.0,
            'sales_amount_with_vat': 0.0
        }

        for col, default in all_columns.items():
            if col not in df_to_save.columns:
                df_to_save[col] = default

        # Числовые колонки
        numeric_cols = [
            'revenue', 'cost_price', 'gross_profit', 'sales_expenses',
            'other_income_expenses', 'net_profit', 'vat_deductible', 'vat_to_budget',
            'purchase_amount_with_vat', 'sales_amount_without_vat', 'sales_amount_with_vat'
        ]
        for col in numeric_cols:
            if col in df_to_save.columns:
                df_to_save[col] = pd.to_numeric(df_to_save[col], errors='coerce').fillna(0.0)

        if 'quantity' in df_to_save.columns:
            df_to_save['quantity'] = pd.to_numeric(df_to_save['quantity'], errors='coerce').fillna(0).astype(int)

        if 'id' in df_to_save.columns:
            df_to_save = df_to_save.drop(columns=['id'])

        df_to_save.to_sql('reports', self.conn, if_exists='append', index=False)
        self.conn.commit()
        return len(df_to_save)

    def get_all_data(self):
        query = "SELECT * FROM reports ORDER BY period_start DESC, company"
        return pd.read_sql_query(query, self.conn)

    def get_filtered_data(self, company=None, date_from=None, date_to=None, product_group=None, doc_type=None):
        query = "SELECT * FROM reports WHERE 1=1"
        params = []

        if company and company != "Все компании":
            query += " AND company = ?"
            params.append(company)

        if date_from:
            query += " AND period_start >= ?"
            params.append(date_from)

        if date_to:
            query += " AND period_end <= ?"
            params.append(date_to)

        if product_group and product_group != "Все группы":
            query += " AND product_group = ?"
            params.append(product_group)

        if doc_type:
            query += " AND doc_type = ?"
            params.append(doc_type)

        query += " ORDER BY period_start DESC, company"
        return pd.read_sql_query(query, self.conn, params=params)


# ==================== ГЛАВНОЕ ОКНО ====================
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.db = DatabaseManager()
        self.current_df = None
        self.settings = QSettings("DeerTuund", "BuhTuundOtchet")
        
        # Пути из настроек
        self.load_folder = self.settings.value("load_folder", "")
        self.save_folder = self.settings.value("save_folder", "")
        self.db_load_folder = self.settings.value("db_load_folder", "")
        self.db_save_folder = self.settings.value("db_save_folder", "")
        
        self.init_ui()
        self.load_last_database()
        self.load_last_folder()         # загружаем последнюю папку

    # ==================== НАСТРОЙКИ ====================
    def load_settings(self):
        self.load_folder = self.settings.value("load_folder", "")
        self.save_folder = self.settings.value("save_folder", "")
        self.db_load_folder = self.settings.value("db_load_folder", "")
        self.db_save_folder = self.settings.value("db_save_folder", "")

    def save_settings(self):
        self.settings.setValue("load_folder", self.load_folder)
        self.settings.setValue("save_folder", self.save_folder)
        self.settings.setValue("db_load_folder", self.db_load_folder)
        self.settings.setValue("db_save_folder", self.db_save_folder)

    def show_settings(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Настройки")
        dialog.setModal(True)
        layout = QVBoxLayout(dialog)

        # Папка загрузки данных
        load_layout = QHBoxLayout()
        load_layout.addWidget(QLabel("Папка загрузки данных:"))
        self.load_folder_edit = QLineEdit(self.load_folder)
        load_layout.addWidget(self.load_folder_edit)
        load_btn = QPushButton("Обзор...")
        load_btn.clicked.connect(lambda: self._choose_folder(self.load_folder_edit, "load_folder"))
        load_layout.addWidget(load_btn)
        layout.addLayout(load_layout)

        # Папка сохранения отчетов
        save_layout = QHBoxLayout()
        save_layout.addWidget(QLabel("Папка сохранения отчетов:"))
        self.save_folder_edit = QLineEdit(self.save_folder)
        save_layout.addWidget(self.save_folder_edit)
        save_btn = QPushButton("Обзор...")
        save_btn.clicked.connect(lambda: self._choose_folder(self.save_folder_edit, "save_folder"))
        save_layout.addWidget(save_btn)
        layout.addLayout(save_layout)

        # Папка загрузки БД
        db_load_layout = QHBoxLayout()
        db_load_layout.addWidget(QLabel("Папка загрузки БД:"))
        self.db_load_folder_edit = QLineEdit(self.db_load_folder)
        db_load_layout.addWidget(self.db_load_folder_edit)
        db_load_btn = QPushButton("Обзор...")
        db_load_btn.clicked.connect(lambda: self._choose_folder(self.db_load_folder_edit, "db_load_folder"))
        db_load_layout.addWidget(db_load_btn)
        layout.addLayout(db_load_layout)

        # Папка сохранения БД
        db_save_layout = QHBoxLayout()
        db_save_layout.addWidget(QLabel("Папка сохранения БД:"))
        self.db_save_folder_edit = QLineEdit(self.db_save_folder)
        db_save_layout.addWidget(self.db_save_folder_edit)
        db_save_btn = QPushButton("Обзор...")
        db_save_btn.clicked.connect(lambda: self._choose_folder(self.db_save_folder_edit, "db_save_folder"))
        db_save_layout.addWidget(db_save_btn)
        layout.addLayout(db_save_layout)

        btn_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btn_box.accepted.connect(lambda: self._save_settings_from_dialog(dialog))
        btn_box.rejected.connect(dialog.reject)
        layout.addWidget(btn_box)

        dialog.exec()

    def _choose_folder(self, line_edit, setting_key):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку")
        if folder:
            line_edit.setText(folder)
            setattr(self, setting_key, folder)

    def _save_settings_from_dialog(self, dialog):
        self.load_folder = self.load_folder_edit.text()
        self.save_folder = self.save_folder_edit.text()
        self.db_load_folder = self.db_load_folder_edit.text()
        self.db_save_folder = self.db_save_folder_edit.text()
        self.save_settings()
        dialog.accept()
    #==================================================================================
    # ==================== ЗАГРУЗКА ПОСЛЕДНЕЙ БД ====================
    def load_last_database(self):
        """Загружает последнюю использованную базу данных при старте"""
        last_db = self.settings.value("last_database", "")
        
        # Если последняя БД не существует или была удалена
        if not last_db or not os.path.exists(last_db):
            # Создаем новую БД по умолчанию
            self.db = DatabaseManager()
            self.current_df = pd.DataFrame()
            self.display_data(self.current_df)
            self.update_summary()
            self.update_charts()
            self.update_filter_combos()
            print("Создана новая база данных по умолчанию")
            return
        
        try:
            self.db.conn.close()
            self.db = DatabaseManager(db_path=last_db)
            self.current_df = self.db.get_all_data()
            self.display_data(self.current_df)
            self.update_summary()
            self.update_charts()
            self.update_filter_combos()
            print(f"Загружена последняя БД: {last_db}")
        except Exception as e:
            print(f"Не удалось загрузить последнюю БД: {e}")
            # В случае ошибки создаем новую БД
            self.db = DatabaseManager()
            self.current_df = pd.DataFrame()
            self.display_data(self.current_df)
            self.update_summary()
            self.update_charts()
            self.update_filter_combos()

    # ====================================================================================
    # """Загружает последнюю использованную папку в дерево файлов"""
    def load_last_folder(self):
        """Загружает последнюю использованную папку в дерево файлов"""
        last_folder = self.settings.value("last_folder", "")
        if last_folder and os.path.exists(last_folder):
            self.load_folder_tree(last_folder)

    # ==========================================================================================
    # ==================== """Сохраняет настройки при закрытии программы""" ====================
    def closeEvent(self, event):
        """Сохраняет настройки при закрытии программы"""
        # Сохраняем путь к текущей базе данных
        cursor = self.db.conn.execute("PRAGMA database_list")
        row = cursor.fetchone()
        if row and row[2]:  # есть путь к файлу
            self.settings.setValue("last_database", row[2])
        
        # Сохраняем последнюю открытую папку в дереве
        root = self.tree_widget.topLevelItem(0)
        if root:
            folder_path = root.data(0, Qt.ItemDataRole.UserRole)
            if folder_path and os.path.isdir(folder_path):
                self.settings.setValue("last_folder", folder_path)
        
        event.accept()

    # ==================== ИНИЦИАЛИЗАЦИЯ ИНТЕРФЕЙСА ====================
    def init_ui(self):
        self.setWindowTitle("BuhTuundOtchet")
        self.setGeometry(100, 100, 1400, 800)
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QMenuBar {
                background-color: #2c3e50;
                color: white;
                font-weight: bold;
            }
            QMenuBar::item {
                background-color: #2c3e50;
                color: white;
                padding: 5px 10px;
            }
            QMenuBar::item:selected {
                background-color: #3498db;
            }
            QMenu {
                background-color: #ecf0f1;
                border: 1px solid #bdc3c7;
            }
            QMenu::item:selected {
                background-color: #3498db;
                color: white;
            }
            QToolBar {
                background-color: #34495e;
                spacing: 5px;
                padding: 5px;
            }
            QToolButton {
                background-color: #3498db;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 12px;
                font-weight: bold;
            }
            QToolButton:hover {
                background-color: #2980b9;
            }
            QTableView {
                background-color: white;
                alternate-background-color: #f8f9fa;
                selection-background-color: #3498db;
                gridline-color: #dee2e6;
                font-size: 11pt;
            }
            QHeaderView::section {
                background-color: #34495e;
                color: white;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QComboBox, QLineEdit {
                padding: 6px;
                border: 1px solid #bdc3c7;
                border-radius: 4px;
                background-color: white;
            }
            QLabel {
                font-weight: bold;
                color: #2c3e50;
            }
        """)

        self.create_menus()

        # Центральный виджет - ОБЯЗАТЕЛЬНО!!!
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        # Панель фильтров
        filter_layout = QHBoxLayout()
        
        self.company_combo = QComboBox()
        self.company_combo.addItems(["Все компании"])
        
        self.period_combo = QComboBox()
        self.period_combo.addItems(["Все периоды"])
        
        self.group_combo = QComboBox()
        self.group_combo.addItems(["Все группы"])
        
        filter_layout.addWidget(QLabel("Компания:"))
        filter_layout.addWidget(self.company_combo)
        filter_layout.addWidget(QLabel("Период:"))
        filter_layout.addWidget(self.period_combo)
        filter_layout.addWidget(QLabel("Товарная группа:"))
        filter_layout.addWidget(self.group_combo)
        
        self.apply_filter_btn = QPushButton("Применить фильтр")
        self.apply_filter_btn.clicked.connect(self.apply_filters)
        self.apply_filter_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                font-weight: bold;
                padding: 8px 16px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #219653;
            }
        """)
        filter_layout.addWidget(self.apply_filter_btn)
        
        # ДОБАВЛЯЕМ ФИЛЬТРЫ В MAIN_LAYOUT
        main_layout.addLayout(filter_layout)

        # СОЗДАЁМ СПЛИТТЕР - ОН БУДЕТ РАЗДЕЛЯТЬ ЛЕВУЮ И ПРАВУЮ ПАНЕЛИ
        self.splitter = QSplitter(Qt.Orientation.Horizontal)

        # Левая панель с деревом
        self.left_panel = QWidget()
        left_layout = QVBoxLayout(self.left_panel)
        left_layout.setContentsMargins(2, 2, 2, 2)

        self.select_root_btn = QPushButton("Выбрать папку...")
        self.select_root_btn.clicked.connect(self.choose_root_folder)
        left_layout.addWidget(self.select_root_btn)

        self.tree_widget = QTreeWidget()
        self.tree_widget.setHeaderHidden(True)
        self.tree_widget.setSelectionMode(QAbstractItemView.SelectionMode.NoSelection)
        left_layout.addWidget(self.tree_widget)

        self.tree_widget.itemChanged.connect(self._handle_item_changed)
        self.tree_widget.itemChanged.connect(self._update_process_button_state)

        self.splitter.addWidget(self.left_panel)

        # Правая панель с вкладками
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)

        self.tab_widget = QTabWidget()
        self.tab_widget.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #bdc3c7;
                background-color: white;
            }
            QTabBar::tab {
                background-color: #ecf0f1;
                padding: 10px 20px;
                margin-right: 2px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
            }
            QTabBar::tab:selected {
                background-color: #3498db;
                color: white;
                font-weight: bold;
            }
        """)

        # Вкладка с таблицей
        self.table_tab = QWidget()
        table_layout = QVBoxLayout(self.table_tab)

        self.table_view = QTableView()
        self.table_model = QStandardItemModel()
        self.table_view.setModel(self.table_model)
        self.table_view.setAlternatingRowColors(True)
        self.table_view.setSortingEnabled(True)

        headers = [
            "ID", "Компания", "Период с", "Период по", "Тип", "Группа",
            "Продавец", "Покупатель", "Номенклатура",
            "№ сч/ф", "Дата сч/ф", "Код опер.", "Дата принятия", "Плат. док.",
            "Сумма покупки с НДС", "Сумма продажи без НДС", "Сумма продажи с НДС",
            "Выручка", "Себестоимость", "Валовая прибыль",
            "Расходы на продажу", "Прочие доходы/расходы", "Чистая прибыль",
            "НДС покупки", "НДС продажи", "Кол-во", "Дата импорта"
        ]
        self.table_model.setHorizontalHeaderLabels(headers)

        table_layout.addWidget(self.table_view)

        # Панель итогов
        summary_layout = QHBoxLayout()
        
        self.revenue_with_vat_label = QLabel("Выручка с НДС: 0 ₽")
        self.expenses_with_vat_label = QLabel("Затраты с НДС: 0 ₽")
        self.gross_profit_with_vat_label = QLabel("Валовая прибыль: 0 ₽")
        self.profit_without_vat_label = QLabel("Прибыль без НДС: 0 ₽")
        self.vat_to_budget_net_label = QLabel("НДС в бюджет: 0 ₽")
        self.profit_tax_label = QLabel("Налог на прибыль: 0 ₽")

        for label in [self.revenue_with_vat_label, self.expenses_with_vat_label,
                    self.gross_profit_with_vat_label, self.profit_without_vat_label,
                    self.vat_to_budget_net_label, self.profit_tax_label]:
            label.setStyleSheet("""
                QLabel {
                    background-color: #ecf0f1;
                    padding: 8px 12px;
                    border-radius: 4px;
                    font-weight: bold;
                    color: #2c3e50;
                    border: 1px solid #bdc3c7;
                }
            """)

        summary_layout.addWidget(self.revenue_with_vat_label)
        summary_layout.addWidget(self.expenses_with_vat_label)
        summary_layout.addWidget(self.gross_profit_with_vat_label)
        summary_layout.addWidget(self.profit_without_vat_label)
        summary_layout.addWidget(self.vat_to_budget_net_label)
        summary_layout.addWidget(self.profit_tax_label)
        summary_layout.addStretch()

        table_layout.addLayout(summary_layout)
        
        #-----------------------------------------------------------------------
        # Вкладка с графиками - КАЖДЫЙ ГРАФИК НА ОТДЕЛЬНОЙ СТРОКЕ
        self.charts_tab = QWidget()
        charts_layout = QVBoxLayout(self.charts_tab)

        # Создаём область с прокруткой
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        # Контейнер для всех графиков
        charts_container = QWidget()
        charts_container_layout = QVBoxLayout(charts_container)
        charts_container_layout.setSpacing(30)  # Большие отступы между графиками
        charts_container_layout.setContentsMargins(10, 10, 10, 10)

        # Первый график - отдельно
        self.figure1, self.ax1 = plt.subplots(figsize=(10, 8))  # Увеличенный размер
        self.figure1.patch.set_facecolor('#f5f5f5')
        self.canvas1 = FigureCanvas(self.figure1)
        self.canvas1.setMinimumHeight(500)
        charts_container_layout.addWidget(QLabel("График 1. Распределение прибыли по товарным группам"))
        charts_container_layout.addWidget(self.canvas1)
        charts_container_layout.addSpacing(20)

        # Второй график - отдельно
        self.figure2, self.ax2 = plt.subplots(figsize=(10, 8))
        self.figure2.patch.set_facecolor('#f5f5f5')
        self.canvas2 = FigureCanvas(self.figure2)
        self.canvas2.setMinimumHeight(500)
        charts_container_layout.addWidget(QLabel("График 2. ТОП-5 товаров по прибыльности"))
        charts_container_layout.addWidget(self.canvas2)
        charts_container_layout.addSpacing(20)

        # Третий график - отдельно
        self.figure3, self.ax3 = plt.subplots(figsize=(10, 8))
        self.figure3.patch.set_facecolor('#f5f5f5')
        self.canvas3 = FigureCanvas(self.figure3)
        self.canvas3.setMinimumHeight(500)
        charts_container_layout.addWidget(QLabel("График 3. Закупки с НДС по кварталам"))
        charts_container_layout.addWidget(self.canvas3)
        charts_container_layout.addSpacing(20)

        # Четвертый график - отдельно
        self.figure4, self.ax4 = plt.subplots(figsize=(10, 8))
        self.figure4.patch.set_facecolor('#f5f5f5')
        self.canvas4 = FigureCanvas(self.figure4)
        self.canvas4.setMinimumHeight(500)
        charts_container_layout.addWidget(QLabel("График 4. Выручка с НДС по кварталам"))
        charts_container_layout.addWidget(self.canvas4)
        charts_container_layout.addSpacing(20)

        # Пятый график - отдельно
        self.figure5, self.ax5 = plt.subplots(figsize=(10, 8))
        self.figure5.patch.set_facecolor('#f5f5f5')
        self.canvas5 = FigureCanvas(self.figure5)
        self.canvas5.setMinimumHeight(500)
        charts_container_layout.addWidget(QLabel("График 5. НДС в бюджет по кварталам"))
        charts_container_layout.addWidget(self.canvas5)
        charts_container_layout.addSpacing(20)

        # Шестой график - отдельно
        self.figure6, self.ax6 = plt.subplots(figsize=(10, 8))
        self.figure6.patch.set_facecolor('#f5f5f5')
        self.canvas6 = FigureCanvas(self.figure6)
        self.canvas6.setMinimumHeight(500)
        charts_container_layout.addWidget(QLabel("График 6. НДС по выручке по кварталам"))
        charts_container_layout.addWidget(self.canvas6)
        charts_container_layout.addSpacing(20)

        # Седьмой график - отдельно
        self.figure7, self.ax7 = plt.subplots(figsize=(10, 8))
        self.figure7.patch.set_facecolor('#f5f5f5')
        self.canvas7 = FigureCanvas(self.figure7)
        self.canvas7.setMinimumHeight(500)
        charts_container_layout.addWidget(QLabel("График 7. НДС по затратам по кварталам"))
        charts_container_layout.addWidget(self.canvas7)
        charts_container_layout.addSpacing(20)

        # Восьмой график - отдельно
        self.figure8, self.ax8 = plt.subplots(figsize=(10, 8))
        self.figure8.patch.set_facecolor('#f5f5f5')
        self.canvas8 = FigureCanvas(self.figure8)
        self.canvas8.setMinimumHeight(500)
        charts_container_layout.addWidget(QLabel("График 8. Валовая прибыль по кварталам"))
        charts_container_layout.addWidget(self.canvas8)
        charts_container_layout.addSpacing(20)

        # Девятый график - отдельно
        self.figure9, self.ax9 = plt.subplots(figsize=(10, 8))
        self.figure9.patch.set_facecolor('#f5f5f5')
        self.canvas9 = FigureCanvas(self.figure9)
        self.canvas9.setMinimumHeight(500)
        charts_container_layout.addWidget(QLabel("График 9. Затраты по кварталам (все налоги и закупки)"))
        charts_container_layout.addWidget(self.canvas9)

        # Кнопка обновления
        charts_btn_layout = QHBoxLayout()
        self.update_charts_btn = QPushButton("Обновить графики")
        self.update_charts_btn.clicked.connect(self.update_charts)
        self.update_charts_btn.setStyleSheet(self.apply_filter_btn.styleSheet())
        charts_container_layout.addLayout(charts_btn_layout)

        charts_container_layout.addStretch()
        scroll_area.setWidget(charts_container)
        charts_layout.addWidget(scroll_area)

        self.tab_widget.addTab(self.table_tab, "📊 Таблица данных")
        self.tab_widget.addTab(self.charts_tab, "📈 Графики и анализ")

        right_layout.addWidget(self.tab_widget)
        self.splitter.addWidget(right_panel)
        self.splitter.setSizes([250, self.width() - 250])

        # !!! ВАЖНО - ДОБАВЛЯЕМ СПЛИТТЕР В MAIN_LAYOUT !!!
        main_layout.addWidget(self.splitter)

        # Инициализация данными
        self.current_df = pd.DataFrame()
        self.display_data(self.current_df)
        self.update_summary()
        self.update_charts()
        self.update_filter_combos()

    #================================================================================
    # Создание меню
    def create_menus(self):
        menubar = self.menuBar()

        # КНОПКА ОБРАБОТАТЬ - ПЕРВАЯ В СТРОКЕ МЕНЮ
        self.process_selected_btn = QPushButton("ОБРАБОТАТЬ")
        self.process_selected_btn.setEnabled(False)
        self.process_selected_btn.setStyleSheet("""
            QPushButton {
                background-color: #ff4444;
                color: white;
                font-weight: bold;
                font-size: 14px;
                padding: 5px 15px;
                border-radius: 4px;
                margin: 2px 5px;
            }
            QPushButton:hover {
                background-color: #ff6666;
            }
            QPushButton:disabled {
                background-color: #cccccc;
                color: #666666;
            }
        """)
        self.process_selected_btn.clicked.connect(self.process_selected_files)
        menubar.setCornerWidget(self.process_selected_btn, Qt.Corner.TopLeftCorner)

        # Меню "База данных"
        db_menu = menubar.addMenu("База данных")
        load_db_action = QAction("Загрузить БД", self)
        load_db_action.triggered.connect(self.load_database)
        db_menu.addAction(load_db_action)

        save_db_action = QAction("Сохранить БД", self)
        save_db_action.triggered.connect(self.save_database)
        db_menu.addAction(save_db_action)

        save_as_action = QAction("Сохранить БД как...", self)
        save_as_action.triggered.connect(self.save_database_as)
        db_menu.addAction(save_as_action)

        clear_db_action = QAction("Очистить БД", self)
        clear_db_action.triggered.connect(self.clear_database)
        db_menu.addAction(clear_db_action)

        export_db_action = QAction("Экспорт БД в Excel", self)
        export_db_action.triggered.connect(self.export_to_excel)
        db_menu.addAction(export_db_action)

        import_template_action = QAction("Импорт из шаблона", self)
        import_template_action.triggered.connect(self.import_from_template)
        db_menu.addAction(import_template_action)

        # Меню "Отчеты"
        report_menu = menubar.addMenu("Отчеты")
        quick_report_action = QAction("Быстрый отчет", self)
        quick_report_action.triggered.connect(self.generate_quick_report)
        report_menu.addAction(quick_report_action)

        report_menu.addSeparator()

        pdf_action = QAction("Экспорт в PDF", self)
        pdf_action.triggered.connect(self.export_to_pdf)
        report_menu.addAction(pdf_action)

        word_action = QAction("Экспорт в Word", self)
        word_action.triggered.connect(self.export_to_word)
        report_menu.addAction(word_action)

        # Меню "Настройки"
        settings_menu = menubar.addMenu("Настройки")
        settings_action = QAction("Настройки программы", self)
        settings_action.triggered.connect(self.show_settings)
        settings_menu.addAction(settings_action)

        # Меню "О программе"
        about_menu = menubar.addMenu("Помощь")
        about_action = QAction("О программе", self)
        about_action.triggered.connect(self.show_about)
        about_menu.addAction(about_action)
       
    #=======================================================
    #  Метод активации кнопки Обработать
    def _update_process_button_state(self):
        files = self.get_checked_files()
        self.process_selected_btn.setEnabled(len(files) > 0)


       # ==================== РАБОТА С ДЕРЕВОМ ФАЙЛОВ ====================
    def choose_root_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку загрузки", self.load_folder)
        if folder:
            self.settings.setValue("input_folder", folder)
            self.load_folder = folder
            self.load_folder_tree(folder)

    # ================================================================================
    # """Загружает дерево файлов из папки и сохраняет путь в настройки"""
    def load_folder_tree(self, folder_path):
        """Загружает дерево файлов из папки и сохраняет путь в настройки"""
        self.tree_widget.clear()
        root_item = QTreeWidgetItem([os.path.basename(folder_path)])
        root_item.setData(0, Qt.ItemDataRole.UserRole, folder_path)
        root_item.setFlags(root_item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
        root_item.setCheckState(0, Qt.CheckState.Unchecked)
        self.tree_widget.addTopLevelItem(root_item)
        self._add_folder_contents(folder_path, root_item)
        root_item.setExpanded(True)
        # Сохраняем путь к папке
        self.settings.setValue("last_folder", folder_path)
    
    # ================================================================================
    def _add_folder_contents(self, path, parent_item):
        try:
            for item in sorted(os.listdir(path)):
                full_path = os.path.join(path, item)
                if os.path.isdir(full_path):
                    child = QTreeWidgetItem([item])
                    child.setData(0, Qt.ItemDataRole.UserRole, full_path)
                    child.setFlags(child.flags() | Qt.ItemFlag.ItemIsUserCheckable)
                    child.setCheckState(0, Qt.CheckState.Unchecked)
                    parent_item.addChild(child)
                    self._add_folder_contents(full_path, child)
                elif item.lower().endswith(('.xlsx', '.xls')):
                    child = QTreeWidgetItem([item])
                    child.setData(0, Qt.ItemDataRole.UserRole, full_path)
                    child.setFlags(child.flags() | Qt.ItemFlag.ItemIsUserCheckable)
                    child.setCheckState(0, Qt.CheckState.Unchecked)
                    parent_item.addChild(child)
        except Exception as e:
            print(f"Ошибка чтения папки {path}: {e}")

    def get_checked_files(self, item=None, files=None):
        if files is None:
            files = []
            root = self.tree_widget.topLevelItem(0)
            if root is None:
                return files
            self.get_checked_files(root, files)
            return files

        if item.checkState(0) == Qt.CheckState.Checked:
            file_path = item.data(0, Qt.ItemDataRole.UserRole)
            if file_path and os.path.isfile(file_path) and not os.path.basename(file_path).startswith('~$'):
                files.append(file_path)
        elif item.checkState(0) == Qt.CheckState.Checked and os.path.isdir(item.data(0, Qt.ItemDataRole.UserRole)):
            folder = item.data(0, Qt.ItemDataRole.UserRole)
            for root, dirs, files_in_folder in os.walk(folder):
                for f in files_in_folder:
                    if f.lower().endswith(('.xlsx', '.xls')) and not f.startswith('~$'):
                        files.append(os.path.join(root, f))
            return

        for i in range(item.childCount()):
            self.get_checked_files(item.child(i), files)

    def process_selected_files(self):
        files = self.get_checked_files()
        if not files:
            QMessageBox.information(self, "Ничего не выбрано", "Не выбрано ни одного файла для обработки.")
            return
        self.process_files(files)

    def _handle_item_changed(self, item, column):
        self.tree_widget.blockSignals(True)
        state = item.checkState(0)
        self._set_children_checkstate(item, state)
        self.tree_widget.blockSignals(False)

    def _set_children_checkstate(self, item, state):
        for i in range(item.childCount()):
            child = item.child(i)
            child.setCheckState(0, state)
            self._set_children_checkstate(child, state)
    #================================================================================
    # ==================== РАБОТА С БАЗОЙ ДАННЫХ ====================
    def clear_database(self):
        reply = QMessageBox.question(self, "Подтверждение",
                                    "Вы действительно хотите удалить все данные из базы?",
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            cursor = self.db.conn.cursor()
            cursor.execute("DELETE FROM reports")
            cursor.execute("DELETE FROM sqlite_sequence WHERE name='reports'")
            self.db.conn.commit()
            self.current_df = pd.DataFrame()
            self.display_data(self.current_df)
            self.update_summary()
            self.update_charts()
            self.update_filter_combos()
            
            # При очистке базы удаляем запись о последней БД (будет создана новая при закрытии)
            self.settings.remove("last_database")
            
            QMessageBox.information(self, "Готово", "База данных очищена")

    # ===============================================================================
    # """Загружает базу данных из выбранного файла .db."""
    def load_database(self):
        """Загружает базу данных из выбранного файла .db."""
        start_dir = self.db_load_folder if self.db_load_folder else ""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите файл базы данных",
            start_dir,
            "SQLite DB (*.db)"
        )
        if not file_path:
            return

        try:
            self.db.conn.close()
            self.db = DatabaseManager(db_path=file_path)
            self.current_df = self.db.get_all_data()
            self.display_data(self.current_df)
            self.update_summary()
            self.update_charts()
            self.update_filter_combos()
            # Сохраняем путь как последнюю БД
            self.settings.setValue("last_database", file_path)
            QMessageBox.information(self, "Успех", f"База данных загружена из {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить базу данных:\n{str(e)}")

    def save_database(self):
        QMessageBox.information(self, "Сохранение", "Все изменения уже сохранены в текущей базе данных.")

    #===================================================================
    # ====================== Сохранить БД как.... ==========================
    def save_database_as(self):
        cursor = self.db.conn.execute("PRAGMA database_list")
        row = cursor.fetchone()
        if row is None or not row[2]:
            QMessageBox.warning(self, "Предупреждение", "Не удалось определить путь к текущей базе данных")
            return
        current_db_path = row[2]

        start_dir = self.db_save_folder if self.db_save_folder else ""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить базу данных как",
            os.path.join(start_dir, "database.db"),
            "SQLite DB (*.db)"
        )
        if not file_path:
            return

        try:
            shutil.copy2(current_db_path, file_path)
            # Сохраняем путь как последнюю БД
            self.settings.setValue("last_database", file_path)
            QMessageBox.information(self, "Успех", f"База данных сохранена как {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить базу данных:\n{str(e)}")

    # ==================== ИМПОРТ ИЗ ШАБЛОНА ====================
    def import_from_template(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите файл для импорта",
            self.load_folder,
            "Excel/CSV Files (*.xlsx *.xls *.csv)"
        )
        if not file_path:
            return

        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path, encoding='utf-8-sig')
            else:
                df = pd.read_excel(file_path)

            if self._map_columns_and_import(df):
                QMessageBox.information(self, "Успех", "Данные импортированы")
                self.current_df = self.db.get_all_data()
                self.display_data(self.current_df)
                self.update_summary()
                self.update_charts()
                self.update_filter_combos()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось импортировать файл:\n{str(e)}")

    def _map_columns_and_import(self, df):
        cursor = self.db.conn.execute("PRAGMA table_info(reports)")
        db_columns = [col[1] for col in cursor.fetchall() if col[1] not in ['id', 'import_date']]

        dialog = QDialog(self)
        dialog.setWindowTitle("Сопоставление колонок")
        dialog.setMinimumWidth(600)
        layout = QVBoxLayout(dialog)

        layout.addWidget(QLabel("Сопоставьте колонки из файла с полями базы данных:"))

        mapping_table = QTableWidget(len(df.columns), 2)
        mapping_table.setHorizontalHeaderLabels(["Колонка в файле", "Поле в БД"])

        combo_boxes = []
        for i, col in enumerate(df.columns):
            mapping_table.setItem(i, 0, QTableWidgetItem(str(col)))
            combo = QComboBox()
            combo.addItems(db_columns)
            combo_boxes.append(combo)
            mapping_table.setCellWidget(i, 1, combo)

        layout.addWidget(mapping_table)

        btn_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btn_box.accepted.connect(dialog.accept)
        btn_box.rejected.connect(dialog.reject)
        layout.addWidget(btn_box)

        if dialog.exec() != QDialog.DialogCode.Accepted:
            return False

        mapping = {}
        for i, combo in enumerate(combo_boxes):
            target_col = combo.currentText()
            if target_col:
                mapping[df.columns[i]] = target_col

        if not mapping:
            QMessageBox.warning(self, "Предупреждение", "Не выбрано ни одного сопоставления")
            return False

        df_import = df.rename(columns=mapping)
        df_import = df_import[[col for col in mapping.values() if col in db_columns]]

        for col in db_columns:
            if col not in df_import.columns:
                df_import[col] = '' if 'date' in col or 'name' in col else 0

        self.db.save_data(df_import)
        return True

    # ==================== ОБРАБОТКА ФАЙЛОВ ====================
    def process_files(self, file_paths):
        total = len(file_paths)
        if total == 0:
            return

        progress = QProgressDialog("Загрузка файлов...", "Отмена", 0, total, self)
        progress.setWindowModality(Qt.WindowModality.WindowModal)

        success_count = 0
        error_files = []

        for i, file_path in enumerate(file_paths):
            if progress.wasCanceled():
                break
            progress.setValue(i)
            progress.setLabelText(f"Обработка: {os.path.basename(file_path)}")

            try:
                saved = self._import_excel_file(file_path)
                if saved > 0:
                    success_count += 1
            except Exception as e:
                error_files.append(f"{os.path.basename(file_path)}: {str(e)}")

        progress.setValue(total)

        if success_count > 0:
            self.current_df = self.db.get_all_data()
            self.display_data(self.current_df)
            self.update_summary()
            self.update_charts()
            self.update_filter_combos()
            print(f"Загружено записей из БД: {len(self.current_df)}")

        msg = f"Успешно загружено: {success_count} из {total}"
        if error_files:
            msg += "\n\nОшибки:\n" + "\n".join(error_files[:5])
            if len(error_files) > 5:
                msg += f"\n... и ещё {len(error_files)-5} ошибок"
        QMessageBox.information(self, "Результат загрузки", msg)

    def _import_excel_file(self, file_path):
        print(f"Обработка файла: {os.path.basename(file_path)}")

        if os.path.basename(file_path).startswith('~$'):
            print("Пропуск временного файла")
            return 0

        if file_path.lower().endswith('.xls') and not file_path.lower().endswith('.xlsx'):
            try:
                import xlrd
            except ImportError:
                raise ImportError("Для чтения файлов .xls установите xlrd: pip install xlrd")

        try:
            df_preview = pd.read_excel(file_path, nrows=30, header=None, dtype=str)
        except Exception as e:
            print(f"Ошибка чтения с dtype=str: {e}")
            try:
                df_preview = pd.read_excel(file_path, nrows=30, header=None)
                df_preview = df_preview.astype(str)
            except Exception as e2:
                print(f"Не удалось прочитать файл {file_path}: {e2}")
                raise ValueError(f"Не удалось прочитать файл: {e2}")

        df_preview = df_preview.fillna('')
        preview_text = ' '.join(df_preview.values.flatten()).lower()
        preview_text = re.sub(r'\s+', ' ', preview_text)

        print(f"preview_text (первые 200): {preview_text[:200]}")

        if 'книга покупок' in preview_text:
            print("-> Распознана книга покупок")
            df = self._parse_purchase_book(file_path)
            return self.db.save_data(df)

        if 'книга продаж' in preview_text:
            print("-> Распознана книга продаж")
            df = self._parse_sales_book(file_path)
            return self.db.save_data(df)

        if 'оборотно-сальдовая ведомость по счету 19' in preview_text or 'анализ счета 19' in preview_text:
            print("-> Распознан ОСВ 19")
            df = self._parse_osv_19_detailed(file_path)
            return self.db.save_data(df)

        if 'оборотно-сальдовая ведомость по счету 41' in preview_text:
            print("-> Распознан ОСВ 41")
            df = self._parse_osv_41_detailed(file_path)
            return self.db.save_data(df)

        if 'оборотно-сальдовая ведомость по счету 44' in preview_text:
            print("-> Распознан ОСВ 44")
            df = self._parse_osv_44_detailed(file_path)
            return self.db.save_data(df)

        if 'оборотно-сальдовая ведомость по счету 60' in preview_text:
            print("-> Распознан ОСВ 60")
            df = self._parse_osv_60_detailed(file_path)
            return self.db.save_data(df)

        if 'оборотно-сальдовая ведомость по счету 62' in preview_text:
            print("-> Распознан ОСВ 62")
            df = self._parse_osv_62_detailed(file_path)
            return self.db.save_data(df)

        if 'оборотно-сальдовая ведомость по счету 68' in preview_text:
            print("-> Распознан ОСВ 68")
            df = self._parse_osv_68_detailed(file_path)
            return self.db.save_data(df)

        if 'оборотно-сальдовая ведомость по счету 90' in preview_text:
            print("-> Распознан ОСВ 90")
            df = self._parse_osv_90_detailed(file_path)
            return self.db.save_data(df)

        if 'оборотно-сальдовая ведомость по счету 91' in preview_text:
            print("-> Распознан ОСВ 91")
            df = self._parse_osv_91_detailed(file_path)
            return self.db.save_data(df)

        if 'отчет по продажам' in preview_text:
            print("-> Распознан отчет по продажам")
            return self._parse_sales_report_detailed(file_path)

        print("-> Не распознан тип, пробуем legacy импорт")
        return self._import_legacy(file_path)

    # ==================== ПАРСЕРЫ ====================
    def _extract_company_by_keyword(self, df, keyword):
        for i in range(min(15, len(df))):
            row = df.iloc[i].tolist()
            for j, cell in enumerate(row):
                if keyword.lower() in cell.lower():
                    for k in range(j+1, len(row)):
                        if row[k].strip():
                            return row[k].strip()
                    break
        return "Неизвестная компания"

    def _extract_base_number(self, cell):
        import re
        if not isinstance(cell, str):
            cell = str(cell)
        match = re.match(r'^(\d+)', cell.strip())
        return int(match.group(1)) if match else None

    def _find_header_row_loose(self, df, min_required=5):
        for i in range(len(df)):
            row = df.iloc[i].tolist()
            expected = 1
            indices = {}
            for col_idx, cell in enumerate(row):
                base = self._extract_base_number(cell)
                if base is not None:
                    if base == expected:
                        indices[base] = col_idx
                        expected += 1
            if expected - 1 >= min_required:
                for col_idx, cell in enumerate(row):
                    base = self._extract_base_number(cell)
                    if base is not None and base not in indices:
                        indices[base] = col_idx
                return i, indices
        return None, None

    def _find_header_row_fallback(self, df, min_count=5):
        best_row = None
        best_indices = {}
        max_count = 0
        for i in range(len(df)):
            row = df.iloc[i].tolist()
            indices = {}
            for col_idx, cell in enumerate(row):
                base = self._extract_base_number(cell)
                if base is not None and base not in indices:
                    indices[base] = col_idx
            if len(indices) >= min_count and len(indices) > max_count:
                max_count = len(indices)
                best_row = i
                best_indices = indices
        if best_row is not None:
            return best_row, best_indices
        return None, None

    def _clean_number(self, value):
        if value is None:
            return 0.0
        if isinstance(value, (int, float)):
            return float(value)
        if isinstance(value, bytes):
            try:
                s = value.decode('utf-8')
            except:
                s = str(value)
        else:
            s = str(value)
        s = s.strip().replace(' ', '').replace(',', '.').replace('−', '-').replace('—', '-')
        import re
        s = re.sub(r'[^\d.-]', '', s)
        try:
            return float(s) if s else 0.0
        except:
            return 0.0

    def _month_name_to_number(self, month_name):
        month_names = {
            'янв': '01', 'фев': '02', 'мар': '03', 'апр': '04', 'май': '05', 'июн': '06',
            'июл': '07', 'авг': '08', 'сен': '09', 'окт': '10', 'ноя': '11', 'дек': '12'
        }
        for key, num in month_names.items():
            if key in month_name.lower():
                return num
        return '01'

    # ========== КНИГА ПОКУПОК ==========
    def _parse_purchase_book(self, file_path):
        import pandas as pd
        import re
        from datetime import datetime

        df = pd.read_excel(file_path, header=None, dtype=str)
        df = df.fillna('').astype(str).apply(lambda col: col.str.strip())

        company = "Неизвестная компания"
        for i in range(min(10, len(df))):
            row = df.iloc[i].tolist()
            for j, cell in enumerate(row):
                if 'покупатель' in cell.lower():
                    for k in range(j+1, len(row)):
                        if row[k].strip():
                            company = row[k].strip()
                            break
                    if company != "Неизвестная компания":
                        break
            if company != "Неизвестная компания":
                break
        print(f"Книга покупок: компания = {company}")

        header_text = ' '.join(df.iloc[:20].values.flatten()).lower()
        period_match = re.search(r'с\s+(\d{2}\.\d{2}\.\d{4})\s+по\s+(\d{2}\.\d{2}\.\d{4})', header_text, re.IGNORECASE)
        if not period_match:
            raise ValueError("Не найден период в книге покупок")
        period_start = datetime.strptime(period_match.group(1), "%d.%m.%Y").strftime("%Y-%m-%d")
        period_end = datetime.strptime(period_match.group(2), "%d.%m.%Y").strftime("%Y-%m-%d")

        header_row_idx, num_to_idx = self._find_header_row_loose(df, min_required=5)
        if header_row_idx is None:
            header_row_idx, num_to_idx = self._find_header_row_fallback(df, min_count=8)
            if header_row_idx is None:
                raise ValueError("Не найдена строка с номерами колонок")

        print(f"Книга покупок: строка с номерами на индексе {header_row_idx}")
        print(f"Соответствие базовых номеров колонкам: {num_to_idx}")

        required_nums = [2, 3, 8, 9, 14, 15]
        for num in required_nums:
            if num not in num_to_idx:
                raise ValueError(f"Не найден номер колонки {num}")

        op_col = num_to_idx[2]
        doc_col = num_to_idx[3]
        accept_col = num_to_idx[8]
        seller_col = num_to_idx[9]
        amount_col = num_to_idx[14]
        vat_col = num_to_idx[15]

        records = []
        current_seller = None
        data_start = header_row_idx + 1

        for i in range(data_start, len(df)):
            row = df.iloc[i].tolist()
            first_cell_raw = row[0] if len(row) > 0 else ''
            first_cell = str(first_cell_raw).strip().lower() if first_cell_raw else ''

            if not first_cell:
                continue

            if 'всего по продавцу' in first_cell:
                current_seller = None
                continue

            if first_cell == 'всего' and 'по продавцу' not in first_cell:
                break

            if first_cell.replace('.', '', 1).replace(',', '').isdigit():
                seller = current_seller
                if seller_col < len(row) and row[seller_col] and row[seller_col].strip():
                    seller = row[seller_col].strip()
                    current_seller = seller
                elif not seller:
                    continue

                doc_str = row[doc_col] if doc_col < len(row) else ''
                doc_number = ''
                doc_date = ''
                if doc_str:
                    parts = re.split(r'\s+от\s+', doc_str, maxsplit=1, flags=re.IGNORECASE)
                    if len(parts) == 2:
                        doc_number = parts[0].strip()
                        doc_date = parts[1].strip()
                    else:
                        doc_number = doc_str

                operation_code = row[op_col] if op_col < len(row) else ''
                acceptance_date = row[accept_col] if accept_col < len(row) else ''

                amount = self._clean_number(row[amount_col] if amount_col < len(row) else '0')
                vat = self._clean_number(row[vat_col] if vat_col < len(row) else '0')

                if amount == 0 and vat == 0:
                    continue

                records.append({
                    'company': company,
                    'period_start': period_start,
                    'period_end': period_end,
                    'doc_type': 'purchase_book',
                    'product_group': 'Покупки',
                    'seller': seller,
                    'buyer': '',
                    'document_number': doc_number,
                    'document_date': doc_date,
                    'operation_code': operation_code,
                    'acceptance_date': acceptance_date,
                    'purchase_amount_with_vat': amount,
                    'sales_amount_with_vat': 0.0,
                    'sales_amount_without_vat': 0.0,
                    'vat_deductible': vat,
                    'vat_to_budget': 0.0,
                    'nomenclature': '',
                    'revenue': 0.0,
                    'cost_price': 0.0,
                    'gross_profit': 0.0,
                    'sales_expenses': 0.0,
                    'other_income_expenses': 0.0,
                    'net_profit': 0.0,
                    'payment_document': '',
                    'quantity': 1
                })
            else:
                if not first_cell[0].isdigit():
                    current_seller = row[0].strip()

        if not records:
            raise ValueError("Не найдено записей в книге покупок")

        print(f"Книга покупок: найдено записей — {len(records)}")
        return pd.DataFrame(records)

    # ========== КНИГА ПРОДАЖ ==========
    def _parse_sales_book(self, file_path):
        import pandas as pd
        import re
        from datetime import datetime

        print(f"Парсер продаж: начало обработки {file_path}")
        df = pd.read_excel(file_path, header=None, dtype=str)
        df = df.fillna('').astype(str).apply(lambda col: col.str.strip())
        print(f"Прочитано строк: {len(df)}")

        company = "Неизвестная компания"
        for i in range(min(10, len(df))):
            row = df.iloc[i].tolist()
            for j, cell in enumerate(row):
                if 'продавец' in cell.lower():
                    for k in range(j+1, len(row)):
                        if row[k].strip():
                            company = row[k].strip()
                            break
                    if company != "Неизвестная компания":
                        break
            if company != "Неизвестная компания":
                break
        print(f"Книга продаж: компания = {company}")

        header_text = ' '.join(df.iloc[:20].values.flatten()).lower()
        period_match = re.search(r'с\s+(\d{2}\.\d{2}\.\d{4})\s+по\s+(\d{2}\.\d{2}\.\d{4})', header_text, re.IGNORECASE)
        if not period_match:
            raise ValueError("Не найден период в книге продаж")
        period_start = datetime.strptime(period_match.group(1), "%d.%m.%Y").strftime("%Y-%m-%d")
        period_end = datetime.strptime(period_match.group(2), "%d.%m.%Y").strftime("%Y-%m-%d")
        print(f"Период: {period_start} - {period_end}")

        header_row_idx, num_to_idx = self._find_header_row_loose(df, min_required=5)
        if header_row_idx is None:
            header_row_idx, num_to_idx = self._find_header_row_fallback(df, min_count=8)
            if header_row_idx is None:
                for debug_i in range(min(20, len(df))):
                    print(f"Строка {debug_i}: {df.iloc[debug_i].tolist()}")
                raise ValueError("Не найдена строка с номерами колонок")

        print(f"Книга продаж: строка с номерами на индексе {header_row_idx}")
        print(f"Соответствие базовых номеров колонкам: {num_to_idx}")

        required_nums = [2, 3, 7, 8, 11, 13, 14, 17]
        for num in required_nums:
            if num not in num_to_idx:
                raise ValueError(f"Не найден базовый номер колонки {num}")

        header_row = df.iloc[header_row_idx].tolist()

        op_col = num_to_idx[2]
        doc_col = num_to_idx[3]
        buyer_col = num_to_idx[7]
        inn_col = num_to_idx[8]
        payment_col = num_to_idx[11]
        amount_without_vat_col = num_to_idx[14]
        vat_col = num_to_idx[17]

        base13_start = num_to_idx[13]
        amount_with_vat_col = None
        for offset in range(10):
            if base13_start + offset >= len(header_row):
                break
            cell = header_row[base13_start + offset]
            clean = re.sub(r'\s+', '', cell.lower())
            if '13б' in clean:
                amount_with_vat_col = base13_start + offset
                break
        if amount_with_vat_col is None:
            raise ValueError("Не найдена колонка '13б' (сумма с НДС)")

        print(f"Индексы Excel: операция={op_col}, документ={doc_col}, покупатель={buyer_col}, ИНН={inn_col}, оплата={payment_col}, сумма без НДС={amount_without_vat_col}, НДС={vat_col}, сумма с НДС={amount_with_vat_col}")

        records = []
        current_buyer = None
        data_start = header_row_idx + 1

        for i in range(data_start, len(df)):
            row = df.iloc[i].tolist()
            first_cell_raw = row[0] if len(row) > 0 else ''
            first_cell = str(first_cell_raw).strip().lower() if first_cell_raw else ''

            if not first_cell:
                continue

            if 'всего по покупателю' in first_cell:
                current_buyer = None
                continue

            if first_cell == 'всего' and 'по покупателю' not in first_cell:
                print(f"Достигнута финальная строка 'Всего' на строке {i}")
                break

            if first_cell.replace('.', '', 1).replace(',', '').isdigit():
                buyer = current_buyer
                if buyer_col < len(row) and row[buyer_col] and row[buyer_col].strip():
                    buyer = row[buyer_col].strip()
                    current_buyer = buyer
                elif not buyer:
                    continue

                doc_str = row[doc_col] if doc_col < len(row) else ''
                doc_number = ''
                doc_date = ''
                if doc_str:
                    parts = re.split(r'\s+от\s+', doc_str, maxsplit=1, flags=re.IGNORECASE)
                    if len(parts) == 2:
                        doc_number = parts[0].strip()
                        doc_date = parts[1].strip()
                    else:
                        doc_number = doc_str

                operation_code = row[op_col] if op_col < len(row) else ''
                inn = row[inn_col] if inn_col < len(row) else ''
                payment_doc = row[payment_col] if payment_col < len(row) else ''

                amount_with_vat = self._clean_number(row[amount_with_vat_col] if amount_with_vat_col < len(row) else '0')
                amount_without_vat = self._clean_number(row[amount_without_vat_col] if amount_without_vat_col < len(row) else '0')
                vat = self._clean_number(row[vat_col] if vat_col < len(row) else '0')

                if amount_with_vat == 0 and amount_without_vat == 0 and vat == 0:
                    continue

                records.append({
                    'company': company,
                    'period_start': period_start,
                    'period_end': period_end,
                    'doc_type': 'sales_book',
                    'product_group': 'Продажи',
                    'seller': '',
                    'buyer': buyer,
                    'document_number': doc_number,
                    'document_date': doc_date,
                    'operation_code': operation_code,
                    'payment_document': payment_doc,
                    'sales_amount_with_vat': amount_with_vat,
                    'sales_amount_without_vat': amount_without_vat,
                    'vat_to_budget': vat,
                    'vat_deductible': 0.0,
                    'nomenclature': '',
                    'revenue': amount_without_vat,
                    'cost_price': 0.0,
                    'gross_profit': 0.0,
                    'sales_expenses': 0.0,
                    'other_income_expenses': 0.0,
                    'net_profit': 0.0,
                    'purchase_amount_with_vat': 0.0,
                    'acceptance_date': '',
                    'quantity': 1
                })
            else:
                if not first_cell[0].isdigit():
                    current_buyer = row[0].strip()

        if not records:
            raise ValueError("Не найдено записей в книге продаж")

        print(f"Книга продаж: найдено записей — {len(records)}")
        return pd.DataFrame(records)

    # ========== ОСВ (заглушки) ==========
    def _parse_osv_19_detailed(self, file_path):
        print("ОСВ 19: найдено записей — 0")
        return pd.DataFrame()

    def _parse_osv_41_detailed(self, file_path):
        print("ОСВ 41: найдено записей — 0")
        return pd.DataFrame()

    def _parse_osv_44_detailed(self, file_path):
        print("ОСВ 44: найдено записей — 0")
        return pd.DataFrame()

    def _parse_osv_60_detailed(self, file_path):
        print("ОСВ 60: найдено записей — 0")
        return pd.DataFrame()

    def _parse_osv_62_detailed(self, file_path):
        print("ОСВ 62: найдено записей — 0")
        return pd.DataFrame()

    def _parse_osv_68_detailed(self, file_path):
        print("ОСВ 68: найдено записей — 0")
        return pd.DataFrame()

    def _parse_osv_90_detailed(self, file_path):
        print("ОСВ 90: найдено записей — 0")
        return pd.DataFrame()

    def _parse_osv_91_detailed(self, file_path):
        print("ОСВ 91: найдено записей — 0")
        return pd.DataFrame()

    def _parse_sales_report_detailed(self, file_path):
        print("Отчет по продажам: заглушка")
        return 0

    def _import_legacy(self, file_path):
        print("Legacy импорт: заглушка")
        return 0

    # ==================== ОТОБРАЖЕНИЕ ДАННЫХ ====================
    def display_data(self, df):
        self.table_model.setRowCount(0)

        column_order = [
            'id', 'company', 'period_start', 'period_end', 'doc_type', 'product_group',
            'seller', 'buyer', 'nomenclature',
            'document_number', 'document_date', 'operation_code', 'acceptance_date', 'payment_document',
            'purchase_amount_with_vat', 'sales_amount_without_vat', 'sales_amount_with_vat',
            'revenue', 'cost_price', 'gross_profit',
            'sales_expenses', 'other_income_expenses', 'net_profit',
            'vat_deductible', 'vat_to_budget', 'quantity', 'import_date'
        ]

        ru_headers = {
            'id': 'ID',
            'company': 'Компания',
            'period_start': 'Период с',
            'period_end': 'Период по',
            'doc_type': 'Тип',
            'product_group': 'Группа',
            'seller': 'Продавец',
            'buyer': 'Покупатель',
            'nomenclature': 'Номенклатура',
            'document_number': '№ сч/ф',
            'document_date': 'Дата сч/ф',
            'operation_code': 'Код опер.',
            'acceptance_date': 'Дата принятия',
            'payment_document': 'Плат. док.',
            'purchase_amount_with_vat': 'Сумма покупки с НДС',
            'sales_amount_without_vat': 'Сумма продажи без НДС',
            'sales_amount_with_vat': 'Сумма продажи с НДС',
            'revenue': 'Выручка',
            'cost_price': 'Себестоимость',
            'gross_profit': 'Валовая прибыль',
            'sales_expenses': 'Расходы на продажу',
            'other_income_expenses': 'Прочие доходы/расходы',
            'net_profit': 'Чистая прибыль',
            'vat_deductible': 'НДС покупки',
            'vat_to_budget': 'НДС продажи',
            'quantity': 'Кол-во',
            'import_date': 'Дата импорта'
        }

        headers = [ru_headers.get(col, col) for col in column_order]
        self.table_model.setHorizontalHeaderLabels(headers)

        if df is None or df.empty:
            return

        for _, row in df.iterrows():
            items = []
            for col in column_order:
                value = row[col] if col in row.index else ''
                if col in ['purchase_amount_with_vat', 'sales_amount_without_vat', 'sales_amount_with_vat',
                           'revenue', 'cost_price', 'gross_profit', 'sales_expenses',
                           'other_income_expenses', 'net_profit', 'vat_deductible', 'vat_to_budget']:
                    if isinstance(value, (int, float)):
                        display_value = f"{value:,.2f} ₽".replace(",", " ")
                    else:
                        display_value = str(value)
                elif col == 'quantity':
                    if isinstance(value, (int, float)):
                        display_value = str(int(value))
                    else:
                        display_value = str(value)
                else:
                    display_value = str(value)
                item = QStandardItem(display_value)
                item.setData(value)
                items.append(item)
            self.table_model.appendRow(items)

        self.table_view.resizeColumnsToContents()

    # ==================== ФИЛЬТРЫ ====================
    def update_filter_combos(self):
        current_company = self.company_combo.currentText()
        current_period = self.period_combo.currentText()
        current_group = self.group_combo.currentText()

        self.company_combo.clear()
        self.period_combo.clear()
        self.group_combo.clear()

        self.company_combo.addItem("Все компании")
        self.period_combo.addItem("Все периоды")
        self.group_combo.addItem("Все группы")

        if self.current_df is not None and not self.current_df.empty:
            companies = sorted(self.current_df['company'].dropna().unique())
            self.company_combo.addItems([str(c) for c in companies])

            periods = set()
            for date_str in self.current_df['period_start'].dropna().unique():
                try:
                    dt = datetime.strptime(date_str, "%Y-%m-%d")
                    periods.add(dt.strftime("%m.%Y"))
                except:
                    pass
            self.period_combo.addItems(sorted(periods))

            groups = sorted(self.current_df['product_group'].dropna().unique())
            self.group_combo.addItems([str(g) for g in groups])

        index = self.company_combo.findText(current_company)
        if index >= 0:
            self.company_combo.setCurrentIndex(index)
        index = self.period_combo.findText(current_period)
        if index >= 0:
            self.period_combo.setCurrentIndex(index)
        index = self.group_combo.findText(current_group)
        if index >= 0:
            self.group_combo.setCurrentIndex(index)

    def _period_to_dates(self, period_str):
        import calendar
        try:
            month, year = period_str.split('.')
            month = int(month)
            year = int(year)
            start_date = f"{year:04d}-{month:02d}-01"
            last_day = calendar.monthrange(year, month)[1]
            end_date = f"{year:04d}-{month:02d}-{last_day:02d}"
            return start_date, end_date
        except:
            return None, None

    def apply_filters(self):
        company = self.company_combo.currentText()
        period = self.period_combo.currentText()
        product_group = self.group_combo.currentText()

        date_from = None
        date_to = None
        if period != "Все периоды":
            date_from, date_to = self._period_to_dates(period)

        filtered_df = self.db.get_filtered_data(
            company=company if company != "Все компании" else None,
            date_from=date_from,
            date_to=date_to,
            product_group=product_group if product_group != "Все группы" else None
        )

        if not filtered_df.empty:
            self.current_df = filtered_df
            self.display_data(filtered_df)
            self.update_summary()
            self.update_charts()
        else:
            self.current_df = pd.DataFrame()
            self.display_data(self.current_df)
            self.update_summary()
            self.update_charts()

    # ==================== РАСЧЁТ ФИНАНСОВЫХ ПОКАЗАТЕЛЕЙ ====================
    def calculate_financials(self, df=None):
        if df is None:
            df = self.current_df
        if df is None or df.empty:
            return {
                'revenue_with_vat': 0.0,
                'revenue_without_vat': 0.0,
                'expenses_with_vat': 0.0,
                'expenses_without_vat': 0.0,
                'gross_profit_with_vat': 0.0,
                'profit_without_vat': 0.0,
                'profit_margin': 0.0,
                'vat_sales': 0.0,
                'vat_purchases': 0.0,
                'vat_to_budget_net': 0.0,
                'profit_tax': 0.0,
            }

        sales_df = df[df['doc_type'] == 'sales_book']
        revenue_with_vat = sales_df['sales_amount_with_vat'].sum() if 'sales_amount_with_vat' in sales_df else 0.0
        revenue_without_vat = sales_df['sales_amount_without_vat'].sum() if 'sales_amount_without_vat' in sales_df else 0.0
        vat_sales = sales_df['vat_to_budget'].sum() if 'vat_to_budget' in sales_df else 0.0

        purchases_df = df[df['doc_type'] == 'purchase_book']
        expenses_with_vat = purchases_df['purchase_amount_with_vat'].sum() if 'purchase_amount_with_vat' in purchases_df else 0.0
        vat_purchases = purchases_df['vat_deductible'].sum() if 'vat_deductible' in purchases_df else 0.0
        expenses_without_vat = expenses_with_vat - vat_purchases

        gross_profit_with_vat = revenue_with_vat - expenses_with_vat
        profit_without_vat = revenue_without_vat - expenses_without_vat

        profit_margin = (profit_without_vat / revenue_without_vat * 100) if revenue_without_vat != 0 else 0.0

        vat_to_budget_net = vat_sales - vat_purchases
        profit_tax = profit_without_vat * 0.25  # 25% налог на прибыль

        return {
            'revenue_with_vat': revenue_with_vat,
            'revenue_without_vat': revenue_without_vat,
            'expenses_with_vat': expenses_with_vat,
            'expenses_without_vat': expenses_without_vat,
            'gross_profit_with_vat': gross_profit_with_vat,
            'profit_without_vat': profit_without_vat,
            'profit_margin': profit_margin,
            'vat_sales': vat_sales,
            'vat_purchases': vat_purchases,
            'vat_to_budget_net': vat_to_budget_net,
            'profit_tax': profit_tax,
        }

    def update_summary(self):
        fin = self.calculate_financials()
        self.revenue_with_vat_label.setText(f"Выручка с НДС: {fin['revenue_with_vat']:,.0f} ₽".replace(",", " "))
        self.expenses_with_vat_label.setText(f"Затраты с НДС: {fin['expenses_with_vat']:,.0f} ₽".replace(",", " "))
        self.gross_profit_with_vat_label.setText(f"Валовая прибыль: {fin['gross_profit_with_vat']:,.0f} ₽".replace(",", " "))
        self.profit_without_vat_label.setText(f"Прибыль без НДС: {fin['profit_without_vat']:,.0f} ₽".replace(",", " "))
        self.vat_to_budget_net_label.setText(f"НДС в бюджет: {fin['vat_to_budget_net']:,.0f} ₽".replace(",", " "))
        self.profit_tax_label.setText(f"Налог на прибыль: {fin['profit_tax']:,.0f} ₽".replace(",", " "))

    #===========================================================================================
    # ==================== ГРАФИКИ ====================
    def update_charts(self):
        """Создает 9 отдельных графиков и сохраняет их в файлы"""
        if self.current_df is None or self.current_df.empty:
            # Очищаем все холсты
            for i in range(1, 10):
                canvas = getattr(self, f'canvas{i}', None)
                if canvas:
                    fig = getattr(self, f'figure{i}', None)
                    if fig:
                        for ax in fig.axes:
                            ax.clear()
                            ax.text(0.5, 0.5, 'Нет данных для отображения', 
                                ha='center', va='center', fontsize=12)
                        canvas.draw()
            return

        df_clean = self.current_df.fillna(0)
        
        # Добавляем кварталы
        if 'period_start' in df_clean.columns:
            df_clean['quarter'] = pd.to_datetime(df_clean['period_start']).dt.to_period('Q')
            df_clean['quarter_str'] = df_clean['quarter'].dt.quarter.astype(str) + 'кв'

        sales_df = df_clean[df_clean['doc_type'] == 'sales_book']
        purchases_df = df_clean[df_clean['doc_type'] == 'purchase_book']

        # Словарь для хранения путей к графикам
        self.chart_paths = {}

        # ===== ГРАФИК 1. Распределение прибыли =====
        self.ax1.clear()
        try:
            if 'product_group' in df_clean.columns:
                group_profit = df_clean.groupby('product_group')['net_profit'].sum()
                if not group_profit.empty and group_profit.sum() != 0:
                    colors1 = plt.cm.Set3(np.linspace(0, 1, len(group_profit)))
                    self.ax1.pie(group_profit.values, labels=group_profit.index,
                                autopct='%1.1f%%', colors=colors1, startangle=90)
                    self.ax1.set_title('График 1. Распределение прибыли по товарным группам', fontsize=14)
                else:
                    self.ax1.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
            else:
                self.ax1.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
        except Exception as e:
            self.ax1.text(0.5, 0.5, f'Ошибка', ha='center', va='center')
        
        self.figure1.tight_layout()
        path1 = "temp_chart_1.png"
        self.figure1.savefig(path1, format='png', dpi=150, bbox_inches='tight')
        self.chart_paths['graph1'] = path1
        self.canvas1.draw()

        # ===== ГРАФИК 2. ТОП-5 товаров =====
        self.ax2.clear()
        try:
            if not sales_df.empty and 'nomenclature' in sales_df.columns:
                product_profit = sales_df.groupby('nomenclature')['net_profit'].sum().reset_index()
                product_profit = product_profit[product_profit['nomenclature'] != '']
                if not product_profit.empty:
                    top_products = product_profit.nlargest(5, 'net_profit')
                    labels = [str(x)[:20] + '...' if len(str(x)) > 20 else str(x)
                            for x in top_products['nomenclature']]
                    colors = plt.cm.viridis(np.linspace(0.2, 0.8, len(top_products)))
                    bars = self.ax2.barh(labels, top_products['net_profit'], color=colors)
                    self.ax2.set_title('График 2. ТОП-5 товаров по прибыльности', fontsize=14)
                    self.ax2.set_xlabel('Прибыль, ₽')
                    for bar in bars:
                        width = bar.get_width()
                        if width > 0:
                            self.ax2.text(width, bar.get_y() + bar.get_height()/2,
                                        f'{width:,.0f}'.replace(",", " "),
                                        ha='left', va='center', fontsize=9)
                else:
                    self.ax2.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
            else:
                self.ax2.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
        except Exception as e:
            self.ax2.text(0.5, 0.5, 'Ошибка', ha='center', va='center')
        
        self.figure2.tight_layout()
        path2 = "temp_chart_2.png"
        self.figure2.savefig(path2, format='png', dpi=150, bbox_inches='tight')
        self.chart_paths['graph2'] = path2
        self.canvas2.draw()

        # ===== ГРАФИК 3. Закупки с НДС по кварталам =====
        self.ax3.clear()
        try:
            if not purchases_df.empty and 'quarter_str' in purchases_df.columns:
                purchases_q = purchases_df.groupby('quarter_str')['purchase_amount_with_vat'].sum().reset_index()
                if not purchases_q.empty and purchases_q['purchase_amount_with_vat'].sum() != 0:
                    colors = plt.cm.Oranges(np.linspace(0.3, 0.8, len(purchases_q)))
                    x_pos = range(len(purchases_q))
                    bars = self.ax3.bar(x_pos, purchases_q['purchase_amount_with_vat'], color=colors)
                    self.ax3.set_title('График 3. Закупки с НДС по кварталам', fontsize=14)
                    self.ax3.set_ylabel('Сумма, ₽')
                    self.ax3.set_xticks(x_pos)
                    self.ax3.set_xticklabels(purchases_q['quarter_str'])
                    self.ax3.grid(True, alpha=0.3, axis='y')
                    for bar in bars:
                        height = bar.get_height()
                        if height > 0:
                            self.ax3.text(bar.get_x() + bar.get_width()/2., height,
                                        f'{height:,.0f}'.replace(",", " "),
                                        ha='center', va='bottom', fontsize=9)
                else:
                    self.ax3.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
            else:
                self.ax3.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
        except Exception as e:
            self.ax3.text(0.5, 0.5, 'Ошибка', ha='center', va='center')
        
        self.figure3.tight_layout()
        path3 = "temp_chart_3.png"
        self.figure3.savefig(path3, format='png', dpi=150, bbox_inches='tight')
        self.chart_paths['graph3'] = path3
        self.canvas3.draw()

        # ===== ГРАФИК 4. Выручка с НДС по кварталам =====
        self.ax4.clear()
        try:
            if not sales_df.empty and 'quarter_str' in sales_df.columns:
                revenue_q = sales_df.groupby('quarter_str')['sales_amount_with_vat'].sum().reset_index()
                if not revenue_q.empty and revenue_q['sales_amount_with_vat'].sum() != 0:
                    colors = plt.cm.Blues(np.linspace(0.3, 0.8, len(revenue_q)))
                    x_pos = range(len(revenue_q))
                    bars = self.ax4.bar(x_pos, revenue_q['sales_amount_with_vat'], color=colors)
                    self.ax4.set_title('График 4. Выручка с НДС по кварталам', fontsize=14)
                    self.ax4.set_ylabel('Сумма, ₽')
                    self.ax4.set_xticks(x_pos)
                    self.ax4.set_xticklabels(revenue_q['quarter_str'])
                    self.ax4.grid(True, alpha=0.3, axis='y')
                    for bar in bars:
                        height = bar.get_height()
                        if height > 0:
                            self.ax4.text(bar.get_x() + bar.get_width()/2., height,
                                        f'{height:,.0f}'.replace(",", " "),
                                        ha='center', va='bottom', fontsize=9)
                else:
                    self.ax4.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
            else:
                self.ax4.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
        except Exception as e:
            self.ax4.text(0.5, 0.5, 'Ошибка', ha='center', va='center')
        
        self.figure4.tight_layout()
        path4 = "temp_chart_4.png"
        self.figure4.savefig(path4, format='png', dpi=150, bbox_inches='tight')
        self.chart_paths['graph4'] = path4
        self.canvas4.draw()

        # ===== ГРАФИК 5. НДС в бюджет по кварталам =====
        self.ax5.clear()
        try:
            if 'quarter_str' in df_clean.columns:
                vat_budget = df_clean.groupby('quarter_str').apply(
                    lambda x: x['vat_to_budget'].sum() - x['vat_deductible'].sum()
                ).reset_index(name='vat_budget')
                if not vat_budget.empty and vat_budget['vat_budget'].sum() != 0:
                    colors = plt.cm.Reds(np.linspace(0.3, 0.8, len(vat_budget)))
                    x_pos = range(len(vat_budget))
                    bars = self.ax5.bar(x_pos, vat_budget['vat_budget'], color=colors)
                    self.ax5.set_title('График 5. НДС в бюджет по кварталам', fontsize=14)
                    self.ax5.set_ylabel('Сумма НДС, ₽')
                    self.ax5.set_xticks(x_pos)
                    self.ax5.set_xticklabels(vat_budget['quarter_str'])
                    self.ax5.grid(True, alpha=0.3, axis='y')
                    for bar in bars:
                        height = bar.get_height()
                        if height > 0:
                            self.ax5.text(bar.get_x() + bar.get_width()/2., height,
                                        f'{height:,.0f}'.replace(",", " "),
                                        ha='center', va='bottom', fontsize=9)
                else:
                    self.ax5.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
            else:
                self.ax5.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
        except Exception as e:
            self.ax5.text(0.5, 0.5, 'Ошибка', ha='center', va='center')
        
        self.figure5.tight_layout()
        path5 = "temp_chart_5.png"
        self.figure5.savefig(path5, format='png', dpi=150, bbox_inches='tight')
        self.chart_paths['graph5'] = path5
        self.canvas5.draw()

        # ===== ГРАФИК 6. НДС по выручке по кварталам =====
        self.ax6.clear()
        try:
            if not sales_df.empty and 'quarter_str' in sales_df.columns:
                vat_sales_q = sales_df.groupby('quarter_str')['vat_to_budget'].sum().reset_index()
                if not vat_sales_q.empty and vat_sales_q['vat_to_budget'].sum() != 0:
                    colors = plt.cm.Greens(np.linspace(0.3, 0.8, len(vat_sales_q)))
                    x_pos = range(len(vat_sales_q))
                    bars = self.ax6.bar(x_pos, vat_sales_q['vat_to_budget'], color=colors)
                    self.ax6.set_title('График 6. НДС по выручке по кварталам', fontsize=14)
                    self.ax6.set_ylabel('Сумма НДС, ₽')
                    self.ax6.set_xticks(x_pos)
                    self.ax6.set_xticklabels(vat_sales_q['quarter_str'])
                    self.ax6.grid(True, alpha=0.3, axis='y')
                    for bar in bars:
                        height = bar.get_height()
                        if height > 0:
                            self.ax6.text(bar.get_x() + bar.get_width()/2., height,
                                        f'{height:,.0f}'.replace(",", " "),
                                        ha='center', va='bottom', fontsize=9)
                else:
                    self.ax6.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
            else:
                self.ax6.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
        except Exception as e:
            self.ax6.text(0.5, 0.5, 'Ошибка', ha='center', va='center')
        
        self.figure6.tight_layout()
        path6 = "temp_chart_6.png"
        self.figure6.savefig(path6, format='png', dpi=150, bbox_inches='tight')
        self.chart_paths['graph6'] = path6
        self.canvas6.draw()

        # ===== ГРАФИК 7. НДС по затратам по кварталам =====
        self.ax7.clear()
        try:
            if not purchases_df.empty and 'quarter_str' in purchases_df.columns:
                vat_purchases_q = purchases_df.groupby('quarter_str')['vat_deductible'].sum().reset_index()
                if not vat_purchases_q.empty and vat_purchases_q['vat_deductible'].sum() != 0:
                    colors = plt.cm.Oranges(np.linspace(0.3, 0.8, len(vat_purchases_q)))
                    x_pos = range(len(vat_purchases_q))
                    bars = self.ax7.bar(x_pos, vat_purchases_q['vat_deductible'], color=colors)
                    self.ax7.set_title('График 7. НДС по затратам по кварталам', fontsize=14)
                    self.ax7.set_ylabel('Сумма НДС, ₽')
                    self.ax7.set_xticks(x_pos)
                    self.ax7.set_xticklabels(vat_purchases_q['quarter_str'])
                    self.ax7.grid(True, alpha=0.3, axis='y')
                    for bar in bars:
                        height = bar.get_height()
                        if height > 0:
                            self.ax7.text(bar.get_x() + bar.get_width()/2., height,
                                        f'{height:,.0f}'.replace(",", " "),
                                        ha='center', va='bottom', fontsize=9)
                else:
                    self.ax7.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
            else:
                self.ax7.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
        except Exception as e:
            self.ax7.text(0.5, 0.5, 'Ошибка', ha='center', va='center')
        
        self.figure7.tight_layout()
        path7 = "temp_chart_7.png"
        self.figure7.savefig(path7, format='png', dpi=150, bbox_inches='tight')
        self.chart_paths['graph7'] = path7
        self.canvas7.draw()

        # ===== ГРАФИК 8. Валовая прибыль по кварталам =====
        self.ax8.clear()
        try:
            if not sales_df.empty and not purchases_df.empty:
                revenue_q = sales_df.groupby('quarter_str')['sales_amount_with_vat'].sum().reset_index()
                expenses_q = purchases_df.groupby('quarter_str')['purchase_amount_with_vat'].sum().reset_index()
                profit_q = pd.merge(revenue_q, expenses_q, on='quarter_str', how='outer').fillna(0)
                profit_q['gross_profit'] = profit_q['sales_amount_with_vat'] - profit_q['purchase_amount_with_vat']
                if not profit_q.empty and profit_q['gross_profit'].sum() != 0:
                    colors = plt.cm.Purples(np.linspace(0.3, 0.8, len(profit_q)))
                    x_pos = range(len(profit_q))
                    bars = self.ax8.bar(x_pos, profit_q['gross_profit'], color=colors)
                    self.ax8.set_title('График 8. Валовая прибыль по кварталам', fontsize=14)
                    self.ax8.set_ylabel('Прибыль, ₽')
                    self.ax8.set_xticks(x_pos)
                    self.ax8.set_xticklabels(profit_q['quarter_str'])
                    self.ax8.grid(True, alpha=0.3, axis='y')
                    for bar in bars:
                        height = bar.get_height()
                        if height > 0:
                            self.ax8.text(bar.get_x() + bar.get_width()/2., height,
                                        f'{height:,.0f}'.replace(",", " "),
                                        ha='center', va='bottom', fontsize=9)
                else:
                    self.ax8.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
            else:
                self.ax8.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
        except Exception as e:
            self.ax8.text(0.5, 0.5, 'Ошибка', ha='center', va='center')
        
        self.figure8.tight_layout()
        path8 = "temp_chart_8.png"
        self.figure8.savefig(path8, format='png', dpi=150, bbox_inches='tight')
        self.chart_paths['graph8'] = path8
        self.canvas8.draw()

        # ===== ГРАФИК 9. Затраты по кварталам =====
        self.ax9.clear()
        try:
            if not purchases_df.empty and 'quarter_str' in purchases_df.columns:
                expenses_q = purchases_df.groupby('quarter_str')['purchase_amount_with_vat'].sum().reset_index()
                if not expenses_q.empty and expenses_q['purchase_amount_with_vat'].sum() != 0:
                    colors = plt.cm.Reds(np.linspace(0.3, 0.8, len(expenses_q)))
                    x_pos = range(len(expenses_q))
                    bars = self.ax9.bar(x_pos, expenses_q['purchase_amount_with_vat'], color=colors)
                    self.ax9.set_title('График 9. Затраты по кварталам (все налоги и закупки)', fontsize=14)
                    self.ax9.set_ylabel('Сумма затрат, ₽')
                    self.ax9.set_xticks(x_pos)
                    self.ax9.set_xticklabels(expenses_q['quarter_str'])
                    self.ax9.grid(True, alpha=0.3, axis='y')
                    for bar in bars:
                        height = bar.get_height()
                        if height > 0:
                            self.ax9.text(bar.get_x() + bar.get_width()/2., height,
                                        f'{height:,.0f}'.replace(",", " "),
                                        ha='center', va='bottom', fontsize=9)
                else:
                    self.ax9.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
            else:
                self.ax9.text(0.5, 0.5, 'Нет данных', ha='center', va='center')
        except Exception as e:
            self.ax9.text(0.5, 0.5, 'Ошибка', ha='center', va='center')
        
        self.figure9.tight_layout()
        path9 = "temp_chart_9.png"
        self.figure9.savefig(path9, format='png', dpi=150, bbox_inches='tight')
        self.chart_paths['graph9'] = path9
        self.canvas9.draw()
    
    
    #===============================================================
    # """Открывает папку, содержащую указанный файл"""
    def open_containing_folder(self, file_path):
        """Открывает папку, содержащую указанный файл"""
        try:
            folder_path = os.path.dirname(file_path)
            if os.path.exists(folder_path):
                if sys.platform == 'win32':
                    os.startfile(folder_path)
                elif sys.platform == 'darwin':  # macOS
                    import subprocess
                    subprocess.run(['open', folder_path])
                else:  # Linux
                    import subprocess
                    subprocess.run(['xdg-open', folder_path])
        except Exception as e:
            print(f"Не удалось открыть папку: {e}")



    #=====================================================================
    # ==================== ЭКСПОРТ В EXCEL ====================
    def export_to_excel(self):
        if self.current_df is None or self.current_df.empty:
            QMessageBox.warning(self, "Предупреждение", "Нет данных для экспорта")
            return

        ru_headers = {
            'id': 'ID',
            'company': 'Компания',
            'period_start': 'Период с',
            'period_end': 'Период по',
            'doc_type': 'Тип',
            'product_group': 'Группа',
            'seller': 'Продавец',
            'buyer': 'Покупатель',
            'nomenclature': 'Номенклатура',
            'document_number': '№ сч/ф',
            'document_date': 'Дата сч/ф',
            'operation_code': 'Код опер.',
            'acceptance_date': 'Дата принятия',
            'payment_document': 'Плат. док.',
            'purchase_amount_with_vat': 'Сумма покупки с НДС',
            'sales_amount_without_vat': 'Сумма продажи без НДС',
            'sales_amount_with_vat': 'Сумма продажи с НДС',
            'revenue': 'Выручка',
            'cost_price': 'Себестоимость',
            'gross_profit': 'Валовая прибыль',
            'sales_expenses': 'Расходы на продажу',
            'other_income_expenses': 'Прочие доходы/расходы',
            'net_profit': 'Чистая прибыль',
            'vat_deductible': 'НДС покупки',
            'vat_to_budget': 'НДС продажи',
            'quantity': 'Кол-во',
            'import_date': 'Дата импорта'
        }

        df_export = self.current_df.copy()
        df_export.rename(columns=ru_headers, inplace=True)

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить как Excel", 
            os.path.join(self.save_folder, "отчет_buh_tuund.xlsx") if self.save_folder else "отчет_buh_tuund.xlsx",
            "Excel Files (*.xlsx)"
        )
        if not file_path:
            return

        try:
            buf = io.BytesIO()
            self.figure.savefig(buf, format='png', dpi=100, bbox_inches='tight')
            buf.seek(0)

            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                df_export.to_excel(writer, sheet_name='Данные', index=False)

                summary_df = pd.DataFrame({
                    'Показатель': ['Общая выручка', 'НДС продажи', 'НДС покупки', 'НДС в бюджет',
                                   'Валовая прибыль', 'Прибыль без НДС', 'Налог на прибыль',
                                   'Количество записей', 'Дата экспорта'],
                    'Значение': [
                        f"{self.current_df['revenue'].sum():,.0f} ₽".replace(",", " "),
                        f"{self.current_df['vat_to_budget'].sum():,.0f} ₽".replace(",", " "),
                        f"{self.current_df['vat_deductible'].sum():,.0f} ₽".replace(",", " "),
                        f"{self.current_df['vat_to_budget'].sum() - self.current_df['vat_deductible'].sum():,.0f} ₽".replace(",", " "),
                        f"{self.current_df['gross_profit'].sum():,.0f} ₽".replace(",", " "),
                        f"{self.current_df['net_profit'].sum():,.0f} ₽".replace(",", " "),
                        f"{self.current_df['net_profit'].sum() * 0.25:,.0f} ₽".replace(",", " "),
                        len(self.current_df),
                        datetime.now().strftime("%d.%m.%Y %H:%M")
                    ]
                })
                summary_df.to_excel(writer, sheet_name='Итоги', index=False)

                workbook = writer.book
                for sheet_name in workbook.sheetnames:
                    worksheet = workbook[sheet_name]
                    for column in worksheet.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except:
                                pass
                        adjusted_width = min(max_length + 2, 50)
                        worksheet.column_dimensions[column_letter].width = adjusted_width
                    for cell in worksheet[1]:
                        cell.font = Font(bold=True)

            QMessageBox.information(self, "Успех", f"Файл сохранен: {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при экспорте: {str(e)}")

    #========================================================================================
    # ==================== ЭКСПОРТ В PDF ====================
    def export_to_pdf(self):
        """print(f"chart_paths в PDF: {self.chart_paths if hasattr(self, 'chart_paths') else 'None'}")
        if hasattr(self, 'chart_paths'):
            for key, path in self.chart_paths.items():
                print(f"{key}: {path} exists: {os.path.exists(path)}") """

        """Экспорт отчета в PDF с отдельными графиками"""
        if self.current_df is None or self.current_df.empty:
            QMessageBox.warning(self, "Предупреждение", "Нет данных для экспорта")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить как PDF",
            os.path.join(self.save_folder, f"отчет_buh_tuund_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf") 
            if self.save_folder else f"отчет_buh_tuund_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
            "PDF Files (*.pdf)"
        )
        if not file_path:
            return

        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            # Регистрация шрифта
            try:
                pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
                font_available = True
            except:
                font_available = False

            doc = SimpleDocTemplate(file_path, pagesize=A4,
                                    leftMargin=2*cm, rightMargin=2*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)
            elements = []
            styles = getSampleStyleSheet()

            if font_available:
                for style_name in styles.byName:
                    styles[style_name].fontName = 'Arial'

            # Стили
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName='Arial' if font_available else styles['Heading1'].fontName,
                fontSize=20,
                alignment=TA_CENTER,
                spaceAfter=20,
                textColor=colors.HexColor('#2c3e50')
            )

            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Heading2'],
                fontName='Arial' if font_available else styles['Heading2'].fontName,
                fontSize=14,
                alignment=TA_LEFT,
                spaceAfter=10,
                textColor=colors.HexColor('#34495e')
            )

            # Получаем название компании
            company_name = "Неизвестная компания"
            if not self.current_df.empty and 'company' in self.current_df.columns:
                unique_companies = self.current_df['company'].dropna().unique()
                if len(unique_companies) > 0:
                    company_name = unique_companies[0]

            # ===== ТИТУЛЬНЫЙ ЛИСТ =====
            elements.append(Paragraph(f"БУХГАЛТЕРСКИЙ ОТЧЕТ", title_style))
            elements.append(Paragraph(f"{company_name}", title_style))
            elements.append(Spacer(1, 10))

            # Информация о периоде
            period_str = "не определен"
            if not self.current_df.empty and 'period_start' in self.current_df.columns and 'period_end' in self.current_df.columns:
                try:
                    start_min = self.current_df['period_start'].min()
                    end_max = self.current_df['period_end'].max()
                    start_dt = datetime.strptime(start_min, "%Y-%m-%d")
                    end_dt = datetime.strptime(end_max, "%Y-%m-%d")
                    period_str = f"с {start_dt.strftime('%d.%m.%Y')} по {end_dt.strftime('%d.%m.%Y')}"
                except:
                    period_str = f"с {start_min} по {end_max}"

            elements.append(Paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles['Normal']))
            elements.append(Paragraph(f"Отчетный период: {period_str}", styles['Normal']))
            elements.append(Spacer(1, 20))
            elements.append(PageBreak())

            # ===== ТАБЛИЦА 1. Финансовые показатели =====
            elements.append(Paragraph("Таблица 1. Основные финансовые показатели", subtitle_style))
            elements.append(Spacer(1, 5))

            fin = self.calculate_financials()
            
            table_data = [
                ['Наименование показателя', 'Значение'],
                ['Выручка с НДС', f"{fin['revenue_with_vat']:,.0f} ₽".replace(",", " ")],
                ['Выручка без НДС', f"{fin['revenue_without_vat']:,.0f} ₽".replace(",", " ")],
                ['Затраты с НДС', f"{fin['expenses_with_vat']:,.0f} ₽".replace(",", " ")],
                ['Затраты без НДС', f"{fin['expenses_without_vat']:,.0f} ₽".replace(",", " ")],
                ['Валовая прибыль (с НДС)', f"{fin['gross_profit_with_vat']:,.0f} ₽".replace(",", " ")],
                ['Прибыль без НДС', f"{fin['profit_without_vat']:,.0f} ₽".replace(",", " ")],
                ['Норма прибыли', f"{fin['profit_margin']:.2f}%"],
                ['НДС продажи', f"{fin['vat_sales']:,.0f} ₽".replace(",", " ")],
                ['НДС покупки', f"{fin['vat_purchases']:,.0f} ₽".replace(",", " ")],
                ['НДС в бюджет', f"{fin['vat_to_budget_net']:,.0f} ₽".replace(",", " ")],
                ['Налог на прибыль (25%)', f"{fin['profit_tax']:,.0f} ₽".replace(",", " ")]
            ]

            table = Table(table_data, colWidths=[250, 150])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (1, 1), (1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Arial' if font_available else 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Arial' if font_available else 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            elements.append(table)
            elements.append(Spacer(1, 20))
            # elements.append(PageBreak())

            # ===== ВСЕ 9 ГРАФИКОВ - ПО 2 НА СТРАНИЦУ =====
            if hasattr(self, 'chart_paths'):
                # Страница 1: Графики 1-2
                elements.append(Paragraph("График 1. Распределение прибыли по товарным группам", subtitle_style))
                if 'graph1' in self.chart_paths and os.path.exists(self.chart_paths['graph1']):
                    elements.append(Image(self.chart_paths['graph1'], width=500, height=350))
                elements.append(Spacer(1, 20))
                
                elements.append(Paragraph("График 2. ТОП-5 товаров по прибыльности", subtitle_style))
                if 'graph2' in self.chart_paths and os.path.exists(self.chart_paths['graph2']):
                    elements.append(Image(self.chart_paths['graph2'], width=500, height=350))
                elements.append(Spacer(1, 20))
                # elements.append(PageBreak())
                
                # Страница 2: Графики 3-4
                elements.append(Paragraph("График 3. Закупки с НДС по кварталам", subtitle_style))
                if 'graph3' in self.chart_paths and os.path.exists(self.chart_paths['graph3']):
                    elements.append(Image(self.chart_paths['graph3'], width=500, height=350))
                elements.append(Spacer(1, 20))
                
                elements.append(Paragraph("График 4. Выручка с НДС по кварталам", subtitle_style))
                if 'graph4' in self.chart_paths and os.path.exists(self.chart_paths['graph4']):
                    elements.append(Image(self.chart_paths['graph4'], width=500, height=350))
                elements.append(Spacer(1, 20))
                # elements.append(PageBreak())
                
                # Страница 3: Графики 5-6
                elements.append(Paragraph("График 5. НДС в бюджет по кварталам", subtitle_style))
                if 'graph5' in self.chart_paths and os.path.exists(self.chart_paths['graph5']):
                    elements.append(Image(self.chart_paths['graph5'], width=500, height=350))
                elements.append(Spacer(1, 20))
                
                elements.append(Paragraph("График 6. НДС по выручке по кварталам", subtitle_style))
                if 'graph6' in self.chart_paths and os.path.exists(self.chart_paths['graph6']):
                    elements.append(Image(self.chart_paths['graph6'], width=500, height=350))
                elements.append(Spacer(1, 20))
                # elements.append(PageBreak())
                
                # Страница 4: Графики 7-8
                elements.append(Paragraph("График 7. НДС по затратам по кварталам", subtitle_style))
                if 'graph7' in self.chart_paths and os.path.exists(self.chart_paths['graph7']):
                    elements.append(Image(self.chart_paths['graph7'], width=500, height=350))
                elements.append(Spacer(1, 20))
                
                elements.append(Paragraph("График 8. Валовая прибыль по кварталам", subtitle_style))
                if 'graph8' in self.chart_paths and os.path.exists(self.chart_paths['graph8']):
                    elements.append(Image(self.chart_paths['graph8'], width=500, height=350))
                elements.append(Spacer(1, 20))
                # elements.append(PageBreak())
                
                # Страница 5: График 9
                elements.append(Paragraph("График 9. Затраты по кварталам (все налоги и закупки)", subtitle_style))
                if 'graph9' in self.chart_paths and os.path.exists(self.chart_paths['graph9']):
                    elements.append(Image(self.chart_paths['graph9'], width=500, height=350))
                elements.append(Spacer(1, 20))
                elements.append(PageBreak())

            # ===== ТАБЛИЦА 2. Детальные данные =====
            elements.append(Paragraph("Таблица 2. Детальные данные (первые 15 записей)", subtitle_style))
            elements.append(Spacer(1, 5))

            # Подготовка данных для таблицы
            table_data = [['Период', 'Компания', 'Контрагент', 'Выручка с НДС', 'НДС', 'Прибыль']]
            for _, row in self.current_df.head(15).iterrows():
                # Контрагент
                counterparty = str(row.get('buyer', '') or row.get('seller', '') or row.get('nomenclature', ''))
                if not counterparty or counterparty == 'nan':
                    counterparty = '—'
                
                # Прибыль
                profit = row.get('net_profit', 0)
                if pd.isna(profit):
                    profit = 0
                
                # Период
                period_str = str(row.get('period_start', ''))
                if period_str and period_str != 'nan':
                    try:
                        dt = datetime.strptime(period_str, "%Y-%m-%d")
                        period_str = dt.strftime("%m.%Y")
                    except:
                        period_str = period_str[:7] if len(period_str) >= 7 else '—'
                else:
                    period_str = '—'
                
                # Компания
                company_str = str(row.get('company', ''))
                if not company_str or company_str == 'nan':
                    company_str = '—'
                
                # Выручка
                revenue_val = row.get('sales_amount_with_vat', 0)
                if pd.isna(revenue_val):
                    revenue_val = 0
                
                # НДС
                vat_val = row.get('vat_to_budget', 0)
                if pd.isna(vat_val):
                    vat_val = 0
                
                table_data.append([
                    period_str[:10],
                    company_str[:20],
                    counterparty[:25],
                    f"{revenue_val:,.0f} ₽".replace(",", " "),
                    f"{vat_val:,.0f} ₽".replace(",", " "),
                    f"{profit:,.0f} ₽".replace(",", " ")
                ])

            table2 = Table(table_data, colWidths=[60, 100, 120, 80, 70, 80])
            table2.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('ALIGN', (3, 1), (5, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Arial' if font_available else 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Arial' if font_available else 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            elements.append(table2)
            elements.append(Spacer(1, 20))
            elements.append(PageBreak())

            # ===== АНАЛИЗ И ВЫВОДЫ =====
            elements.append(Paragraph("Анализ финансового состояния", subtitle_style))
            elements.append(Spacer(1, 10))

            # Генерируем текстовый анализ
            analysis_lines = [
                "На основе предоставленных данных можно сделать следующие выводы:",
                "",
                f"✓ Компания {'работает с прибылью' if fin['profit_without_vat'] > 0 else 'работает в убыток'}. " +
                f"{'Чистая прибыль' if fin['profit_without_vat'] > 0 else 'Убыток'} без НДС составляет {abs(fin['profit_without_vat']):,.0f} ₽.",
                "",
                f"✓ Норма прибыли составляет {fin['profit_margin']:.2f}%. " +
                ("Это хороший показатель." if fin['profit_margin'] > 10 else
                "Это низкий показатель, требуется оптимизация." if fin['profit_margin'] < 5 else
                "Это средний показатель."),
                "",
            ]

            if fin['vat_to_budget_net'] > 0:
                vat_percent = fin['vat_to_budget_net'] / fin['revenue_with_vat'] * 100 if fin['revenue_with_vat'] != 0 else 0
                analysis_lines.append(f"✓ НДС к уплате в бюджет составляет {fin['vat_to_budget_net']:,.0f} ₽. Это {vat_percent:.1f}% от выручки.")
            else:
                analysis_lines.append(f"✓ НДС к возмещению из бюджета составляет {abs(fin['vat_to_budget_net']):,.0f} ₽.")
            
            tax_burden = fin['profit_tax'] / fin['revenue_with_vat'] * 100 if fin['revenue_with_vat'] != 0 else 0
            analysis_lines.append(f"✓ Налоговая нагрузка (налог на прибыль) составляет {tax_burden:.1f}% от выручки.")
            analysis_lines.append("")
            analysis_lines.append("Рекомендации:")
            
            if fin['profit_margin'] < 5:
                analysis_lines.append("• Необходимо проанализировать структуру затрат и найти пути их снижения.")
            if fin['expenses_with_vat'] > fin['revenue_with_vat'] * 0.9:
                analysis_lines.append("• Высокая доля затрат в выручке. Требуется оптимизация.")
            if fin['vat_to_budget_net'] < 0:
                analysis_lines.append("• Сумма НДС к возмещению значительна. Проверьте правильность оформления счетов-фактур.")

            for line in analysis_lines:
                elements.append(Paragraph(line, styles['Normal']))
                elements.append(Spacer(1, 3))

            elements.append(Spacer(1, 20))

            # Подпись
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Italic'],
                fontName='Arial' if font_available else styles['Italic'].fontName,
                fontSize=8,
                alignment=TA_CENTER,
                textColor=colors.grey
            )
            elements.append(Paragraph("Сформировано программой BuhTuundOtchet", footer_style))

            # Генерация PDF
            doc.build(elements)

            # Удаляем временные файлы
            if hasattr(self, 'chart_paths'):
                for path in self.chart_paths.values():
                    try:
                        if os.path.exists(path):
                            os.remove(path)
                    except:
                        pass

            QMessageBox.information(self, "Успех", f"PDF файл сохранен: {file_path}")
            
            # Открываем папку с сохраненным файлом
            self.open_containing_folder(file_path)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при экспорте в PDF: {str(e)}")
    
    #====================================================================
    # ==================== ЭКСПОРТ В WORD ====================
    def export_to_word(self):
        """Экспорт отчета в Word с отдельными графиками"""
        if self.current_df is None or self.current_df.empty:
            QMessageBox.warning(self, "Предупреждение", "Нет данных для экспорта")
            return

        default_filename = f"отчет_buh_tuund_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить как Word",
            os.path.join(self.save_folder, default_filename) if self.save_folder else default_filename,
            "Word Files (*.docx)"
        )
        if not file_path:
            return

        try:
            import docx
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = docx.Document()

            # Настройка стилей
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            title_style = doc.styles['Title']
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(44, 62, 80)

            # Получаем название компании
            company_name = "Неизвестная компания"
            if not self.current_df.empty and 'company' in self.current_df.columns:
                unique_companies = self.current_df['company'].dropna().unique()
                if len(unique_companies) > 0:
                    company_name = unique_companies[0]

            # ===== ТИТУЛЬНЫЙ ЛИСТ =====
            title = doc.add_heading('БУХГАЛТЕРСКИЙ ОТЧЕТ', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            company_heading = doc.add_heading(company_name, level=1)
            company_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()

            # Информация о периоде
            period_str = "не определен"
            if not self.current_df.empty and 'period_start' in self.current_df.columns and 'period_end' in self.current_df.columns:
                try:
                    start_min = self.current_df['period_start'].min()
                    end_max = self.current_df['period_end'].max()
                    start_dt = datetime.strptime(start_min, "%Y-%m-%d")
                    end_dt = datetime.strptime(end_max, "%Y-%m-%d")
                    period_str = f"с {start_dt.strftime('%d.%m.%Y')} по {end_dt.strftime('%d.%m.%Y')}"
                except:
                    period_str = f"с {start_min} по {end_max}"

            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_para.add_run(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
            info_para.add_run(f"Отчетный период: {period_str}")

            doc.add_paragraph()
            doc.add_page_break()

            # ===== ТАБЛИЦА 1. Финансовые показатели =====
            doc.add_heading('Таблица 1. Основные финансовые показатели', level=2)
            
            fin = self.calculate_financials()
            
            table = doc.add_table(rows=12, cols=2)
            table.style = 'LightShading-Accent1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Заголовки
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Наименование показателя'
            hdr_cells[1].text = 'Значение'
            
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)

            # Данные
            data = [
                ('Выручка с НДС', f"{fin['revenue_with_vat']:,.0f} ₽"),
                ('Выручка без НДС', f"{fin['revenue_without_vat']:,.0f} ₽"),
                ('Затраты с НДС', f"{fin['expenses_with_vat']:,.0f} ₽"),
                ('Затраты без НДС', f"{fin['expenses_without_vat']:,.0f} ₽"),
                ('Валовая прибыль (с НДС)', f"{fin['gross_profit_with_vat']:,.0f} ₽"),
                ('Прибыль без НДС', f"{fin['profit_without_vat']:,.0f} ₽"),
                ('Норма прибыли', f"{fin['profit_margin']:.2f}%"),
                ('НДС продажи', f"{fin['vat_sales']:,.0f} ₽"),
                ('НДС покупки', f"{fin['vat_purchases']:,.0f} ₽"),
                ('НДС в бюджет', f"{fin['vat_to_budget_net']:,.0f} ₽"),
                ('Налог на прибыль (25%)', f"{fin['profit_tax']:,.0f} ₽")
            ]

            for i, (label, value) in enumerate(data, 1):
                cells = table.rows[i].cells
                cells[0].text = label
                cells[1].text = value.replace(",", " ")
                for paragraph in cells[1].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph()
            doc.add_page_break()

            #---------------------------------------------------------------------------------------
            # ===== ВСЕ 9 ГРАФИКОВ - по 2 шт НА ОТДЕЛЬНОЙ СТРАНИЦЕ =====
            if hasattr(self, 'chart_paths'):
                # Страница 1: Графики 1-2
                doc.add_heading('График 1. Распределение прибыли по товарным группам', level=2)
                if 'graph1' in self.chart_paths and os.path.exists(self.chart_paths['graph1']):
                    doc.add_picture(self.chart_paths['graph1'], width=Inches(6.5))  # Растянули
                doc.add_paragraph()
                
                doc.add_heading('График 2. ТОП-5 товаров по прибыльности', level=2)
                if 'graph2' in self.chart_paths and os.path.exists(self.chart_paths['graph2']):
                    doc.add_picture(self.chart_paths['graph2'], width=Inches(6.5))
                doc.add_paragraph()
                doc.add_page_break()
                
                # Страница 2: Графики 3-4
                doc.add_heading('График 3. Закупки с НДС по кварталам', level=2)
                if 'graph3' in self.chart_paths and os.path.exists(self.chart_paths['graph3']):
                    doc.add_picture(self.chart_paths['graph3'], width=Inches(6.5))
                doc.add_paragraph()
                
                doc.add_heading('График 4. Выручка с НДС по кварталам', level=2)
                if 'graph4' in self.chart_paths and os.path.exists(self.chart_paths['graph4']):
                    doc.add_picture(self.chart_paths['graph4'], width=Inches(6.5))
                doc.add_paragraph()
                doc.add_page_break()
                
                # Страница 3: Графики 5-6
                doc.add_heading('График 5. НДС в бюджет по кварталам', level=2)
                if 'graph5' in self.chart_paths and os.path.exists(self.chart_paths['graph5']):
                    doc.add_picture(self.chart_paths['graph5'], width=Inches(6.5))
                doc.add_paragraph()
                
                doc.add_heading('График 6. НДС по выручке по кварталам', level=2)
                if 'graph6' in self.chart_paths and os.path.exists(self.chart_paths['graph6']):
                    doc.add_picture(self.chart_paths['graph6'], width=Inches(6.5))
                doc.add_paragraph()
                doc.add_page_break()
                
                # Страница 4: Графики 7-8
                doc.add_heading('График 7. НДС по затратам по кварталам', level=2)
                if 'graph7' in self.chart_paths and os.path.exists(self.chart_paths['graph7']):
                    doc.add_picture(self.chart_paths['graph7'], width=Inches(6.5))
                doc.add_paragraph()
                
                doc.add_heading('График 8. Валовая прибыль по кварталам', level=2)
                if 'graph8' in self.chart_paths and os.path.exists(self.chart_paths['graph8']):
                    doc.add_picture(self.chart_paths['graph8'], width=Inches(6.5))
                doc.add_paragraph()
                doc.add_page_break()
                
                # Страница 5: График 9
                doc.add_heading('График 9. Затраты по кварталам (все налоги и закупки)', level=2)
                if 'graph9' in self.chart_paths and os.path.exists(self.chart_paths['graph9']):
                    doc.add_picture(self.chart_paths['graph9'], width=Inches(6.5))
                doc.add_paragraph()
                doc.add_page_break()

            # ===== ТАБЛИЦА 2. Детальные данные =====
            doc.add_heading('Таблица 2. Детальные данные (первые 15 записей)', level=2)

            table2 = doc.add_table(rows=1, cols=6)
            table2.style = 'LightShading-Accent1'
            
            # Заголовки
            hdr_cells2 = table2.rows[0].cells
            headers = ['Период', 'Компания', 'Контрагент', 'Выручка с НДС', 'НДС', 'Прибыль']
            for i, header in enumerate(headers):
                hdr_cells2[i].text = header
                for paragraph in hdr_cells2[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Данные
            for _, row in self.current_df.head(15).iterrows():
                cells = table2.add_row().cells
                
                # Период
                period_str = row.get('period_start', '')
                if period_str and isinstance(period_str, str) and period_str != 'nan':
                    try:
                        dt = datetime.strptime(period_str, "%Y-%m-%d")
                        period_str = dt.strftime("%m.%Y")
                    except:
                        period_str = period_str[:7] if len(period_str) >= 7 else '-'
                else:
                    period_str = '-'
                cells[0].text = str(period_str)
                
                # Компания
                company_val = row.get('company', '')
                if company_val and company_val != 'nan':
                    cells[1].text = str(company_val)[:20]
                else:
                    cells[1].text = '-'
                
                # Контрагент
                counterparty = row.get('buyer', '') or row.get('seller', '') or row.get('nomenclature', '')
                if counterparty and counterparty != 'nan':
                    cells[2].text = str(counterparty)[:25]
                else:
                    cells[2].text = '-'
                
                # Выручка
                revenue = row.get('sales_amount_with_vat', 0)
                if pd.isna(revenue):
                    revenue = 0
                cells[3].text = f"{revenue:,.0f} ₽".replace(",", " ")
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # НДС
                vat = row.get('vat_to_budget', 0)
                if pd.isna(vat):
                    vat = 0
                cells[4].text = f"{vat:,.0f} ₽".replace(",", " ")
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Прибыль
                profit = row.get('net_profit', 0)
                if pd.isna(profit):
                    profit = 0
                cells[5].text = f"{profit:,.0f} ₽".replace(",", " ")
                cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph()
            doc.add_page_break()

            # ===== АНАЛИЗ И ВЫВОДЫ =====
            doc.add_heading('Анализ финансового состояния', level=2)

            analysis_para = doc.add_paragraph()
            analysis_para.add_run('На основе предоставленных данных можно сделать следующие выводы:\n\n').bold = True

            if fin['profit_without_vat'] > 0:
                analysis_para.add_run(f'✓ Компания работает с прибылью. Чистая прибыль без НДС составляет {fin["profit_without_vat"]:,.0f} ₽.\n'.replace(",", " "))
            else:
                analysis_para.add_run(f'✗ Компания работает в убыток. Убыток без НДС составляет {abs(fin["profit_without_vat"]):,.0f} ₽.\n'.replace(",", " "))

            margin_text = f'✓ Норма прибыли составляет {fin["profit_margin"]:.2f}%. '
            if fin['profit_margin'] > 10:
                margin_text += 'Это хороший показатель.'
            elif fin['profit_margin'] < 5:
                margin_text += 'Это низкий показатель, требуется оптимизация.'
            else:
                margin_text += 'Это средний показатель.'
            analysis_para.add_run(margin_text + '\n')

            if fin['vat_to_budget_net'] > 0:
                vat_percent = fin['vat_to_budget_net'] / fin['revenue_with_vat'] * 100 if fin['revenue_with_vat'] != 0 else 0
                analysis_para.add_run(f'✓ НДС к уплате в бюджет составляет {fin["vat_to_budget_net"]:,.0f} ₽. '.replace(",", " "))
                analysis_para.add_run(f'Это {vat_percent:.1f}% от выручки.\n')
            else:
                analysis_para.add_run(f'✓ НДС к возмещению из бюджета составляет {abs(fin["vat_to_budget_net"]):,.0f} ₽.\n'.replace(",", " "))

            tax_burden = fin['profit_tax'] / fin['revenue_with_vat'] * 100 if fin['revenue_with_vat'] != 0 else 0
            analysis_para.add_run(f'✓ Налоговая нагрузка (налог на прибыль) составляет {tax_burden:.1f}% от выручки.\n\n')

            analysis_para.add_run('Рекомендации:\n').bold = True
            if fin['profit_margin'] < 5:
                analysis_para.add_run('• Необходимо проанализировать структуру затрат и найти пути их снижения.\n')
            if fin['expenses_with_vat'] > fin['revenue_with_vat'] * 0.9:
                analysis_para.add_run('• Высокая доля затрат в выручке. Требуется оптимизация.\n')
            if fin['vat_to_budget_net'] < 0:
                analysis_para.add_run('• Сумма НДС к возмещению значительна. Проверьте правильность оформления счетов-фактур.\n')

            doc.add_paragraph()

            # ===== ПОДПИСЬ =====
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.add_run('Сформировано программой BuhTuundOtchet').italic = True

            # ===== СОХРАНЕНИЕ ДОКУМЕНТА =====
            doc.save(file_path)

            # ===== УДАЛЕНИЕ ВРЕМЕННЫХ ФАЙЛОВ =====
            if hasattr(self, 'chart_paths'):
                for path in self.chart_paths.values():
                    try:
                        if os.path.exists(path):
                            os.remove(path)
                    except:
                        pass

            QMessageBox.information(self, "Успех", f"Word файл сохранен: {file_path}")
            
            # Открываем папку с сохраненным файлом
            self.open_containing_folder(file_path)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при экспорте в Word: {str(e)}")

    #=====================================================================================
    # ==================== БЫСТРЫЙ ОТЧЕТ ====================
    def generate_quick_report(self):
        if self.current_df is None or self.current_df.empty:
            QMessageBox.warning(self, "Предупреждение", "Нет данных для отчета")
            return

        fin = self.calculate_financials()

        company_name = "Неизвестная компания"
        if not self.current_df.empty and 'company' in self.current_df.columns:
            unique_companies = self.current_df['company'].dropna().unique()
            if len(unique_companies) > 0:
                company_name = unique_companies[0]

        period_str = "не определен"
        if not self.current_df.empty and 'period_start' in self.current_df.columns and 'period_end' in self.current_df.columns:
            try:
                start_min = self.current_df['period_start'].min()
                end_max = self.current_df['period_end'].max()
                start_dt = datetime.strptime(start_min, "%Y-%m-%d")
                end_dt = datetime.strptime(end_max, "%Y-%m-%d")
                period_str = f"с {start_dt.strftime('%d.%m.%Y')} по {end_dt.strftime('%d.%m.%Y')}"
            except:
                period_str = f"с {start_min} по {end_max}"

        report_plain = f"""БЫСТРЫЙ ОТЧЕТ BUHTUUNDOTCHET
        Компания: {company_name}
        Период анализа: {period_str}
        Товарная группа: {self.group_combo.currentText()}

        ОСНОВНЫЕ ПОКАЗАТЕЛИ:
        - Выручка с НДС: {fin['revenue_with_vat']:,.0f} ₽
        - Выручка без НДС: {fin['revenue_without_vat']:,.0f} ₽
        - Затраты с НДС: {fin['expenses_with_vat']:,.0f} ₽
        - Затраты без НДС: {fin['expenses_without_vat']:,.0f} ₽
        - Валовая прибыль (с НДС): {fin['gross_profit_with_vat']:,.0f} ₽
        - Прибыль без НДС: {fin['profit_without_vat']:,.0f} ₽
        - Норма прибыли: {fin['profit_margin']:.2f}%
        - НДС продажи: {fin['vat_sales']:,.0f} ₽
        - НДС покупки: {fin['vat_purchases']:,.0f} ₽
        - НДС в бюджет: {fin['vat_to_budget_net']:,.0f} ₽
        - Налог на прибыль: {fin['profit_tax']:,.0f} ₽

        ТОП-5 товаров по прибыльности:
        {self._get_top_products_text()}

        Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}"""
        
        # Добавляем логотип в HTML (если хотите)
        logo_html = ""
        if os.path.exists("logo.png"):
            # Кодируем изображение в base64 для вставки в HTML
            import base64
            with open("logo.png", "rb") as f:
                logo_data = base64.b64encode(f.read()).decode()
            logo_html = f'<img src="data:image/png;base64,{logo_data}" width="100" height="100" style="float:right;">'
        
        report_html = f"""
        {logo_html}
        <h3>БЫСТРЫЙ ОТЧЕТ BUHTUUNDOTCHET</h3>
        <p><b>Компания:</b> {company_name}</p>
        <p><b>Период анализа:</b> {period_str}</p>
        <p><b>Товарная группа:</b> {self.group_combo.currentText()}</p>
        <hr>
        <p><b>ОСНОВНЫЕ ПОКАЗАТЕЛИ:</b></p>
        <p>• Выручка с НДС: <span style='color: #27ae60; font-weight: bold;'>{fin['revenue_with_vat']:,.0f} ₽</span></p>
        <p>• Выручка без НДС: <span style='color: #27ae60; font-weight: bold;'>{fin['revenue_without_vat']:,.0f} ₽</span></p>
        <p>• Затраты с НДС: <span style='color: #e74c3c; font-weight: bold;'>{fin['expenses_with_vat']:,.0f} ₽</span></p>
        <p>• Затраты без НДС: <span style='color: #e74c3c; font-weight: bold;'>{fin['expenses_without_vat']:,.0f} ₽</span></p>
        <p>• Валовая прибыль (с НДС): <span style='color: #3498db; font-weight: bold;'>{fin['gross_profit_with_vat']:,.0f} ₽</span></p>
        <p>• Прибыль без НДС: <span style='color: #3498db; font-weight: bold;'>{fin['profit_without_vat']:,.0f} ₽</span></p>
        <p>• Норма прибыли: <span style='color: #f39c12; font-weight: bold;'>{fin['profit_margin']:.2f}%</span></p>
        <p>• НДС продажи: <span style='color: #9b59b6; font-weight: bold;'>{fin['vat_sales']:,.0f} ₽</span></p>
        <p>• НДС покупки: <span style='color: #9b59b6; font-weight: bold;'>{fin['vat_purchases']:,.0f} ₽</span></p>
        <p>• НДС в бюджет: <span style='color: #e67e22; font-weight: bold;'>{fin['vat_to_budget_net']:,.0f} ₽</span></p>
        <p>• Налог на прибыль: <span style='color: #e67e22; font-weight: bold;'>{fin['profit_tax']:,.0f} ₽</span></p>
        <hr>
        <p><b>ТОП-5 товаров по прибыльности:</b></p>
        <pre>{self._get_top_products_text()}</pre>
        <hr>
        <p><i>Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}</i></p>
        """

        dlg = QDialog(self)
        dlg.setWindowTitle("Быстрый отчет")
        dlg.setMinimumSize(600, 500)
        layout = QVBoxLayout(dlg)

        text_edit = QTextEdit()
        text_edit.setHtml(report_html)
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)

        button_box = QDialogButtonBox()
        btn_copy = QPushButton("Копировать")
        btn_copy.clicked.connect(lambda: self._copy_report_to_clipboard(report_plain))
        btn_save = QPushButton("Сохранить как...")
        btn_save.clicked.connect(lambda: self._save_report_to_txt(report_plain))
        btn_close = QPushButton("Закрыть")
        btn_close.clicked.connect(dlg.accept)

        button_box.addButton(btn_copy, QDialogButtonBox.ButtonRole.ActionRole)
        button_box.addButton(btn_save, QDialogButtonBox.ButtonRole.ActionRole)
        button_box.addButton(btn_close, QDialogButtonBox.ButtonRole.RejectRole)

        layout.addWidget(button_box)
        dlg.exec()

    def _copy_report_to_clipboard(self, text):
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        QMessageBox.information(self, "Готово", "Отчет скопирован в буфер обмена")

    def _save_report_to_txt(self, text):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить отчет как...",
            os.path.join(self.save_folder, "отчет.txt") if self.save_folder else "отчет.txt",
            "Text Files (*.txt)"
        )
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                QMessageBox.information(self, "Успех", f"Отчет сохранен в {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить файл: {str(e)}")

    def _get_top_products_text(self):
        profit_col = 'net_profit'
        if profit_col in self.current_df.columns and not self.current_df.empty:
            top_products = self.current_df.nlargest(5, profit_col)[['nomenclature', profit_col]]
            lines = []
            for _, row in top_products.iterrows():
                name = row['nomenclature'] if row['nomenclature'] else "Без названия"
                lines.append(f"{name}: {row[profit_col]:,.0f} ₽".replace(",", " "))
            return "\n".join(lines)
        return "Нет данных"

    # ==================== О ПРОГРАММЕ ====================
    def show_about(self):
        # Создаём диалог
        about_dialog = QDialog(self)
        about_dialog.setWindowTitle("О программе BuhTuundOtchet")
        about_dialog.setMinimumWidth(500)
        layout = QVBoxLayout(about_dialog)
        
        # Добавляем логотип
        if os.path.exists("logo.png"):
            logo_label = QLabel()
            pixmap = QPixmap("logo.png")
            # Масштабируем до разумного размера (например, 100x100)
            scaled_pixmap = pixmap.scaled(100, 100, Qt.AspectRatioMode.KeepAspectRatio, 
                                        Qt.TransformationMode.SmoothTransformation)
            logo_label.setPixmap(scaled_pixmap)
            logo_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(logo_label)
        
        # Текст о программе
        about_text = """<h2>Программа BuhTuundOtchet</h2>
        <p><b>Версия программы:</b> v7.3.0</p>
        <p><b>Разработчик:</b> Deer Tuund (C) 2026</p>
        <p><b>Для связи:</b> vaspull9@gmail.com</p>
        <hr>
        <p>Программа для разработки отчетов из 1С</p>
        <p><b>Возможности:</b></p>
        <ul>
            <li>Импорт данных из Excel (выгрузок 1С)</li>
            <li>Хранение данных в SQLite базе</li>
            <li>Фильтрация по компаниям, периодам, товарным группам</li>
            <li>Расчет валовой и чистой прибыли</li>
            <li>Расчет НДС продажи, покупки, в бюджет</li>
            <li>Расчет налога на прибыль (25%)</li>
            <li>Визуализация данных (графики и диаграммы)</li>
            <li>Экспорт в Excel, PDF, Word</li>
            <li>Современный интерфейс с меню и настройками</li>
        </ul>
        <p><b>Используемые технологии:</b> Python, PyQt6, Pandas, Matplotlib, SQLite, ReportLab, python-docx</p>
        """
        
        text_label = QLabel(about_text)
        text_label.setTextFormat(Qt.TextFormat.RichText)
        text_label.setWordWrap(True)
        layout.addWidget(text_label)
        
        # Кнопка закрытия
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok)
        button_box.accepted.connect(about_dialog.accept)
        layout.addWidget(button_box)
        
        about_dialog.exec()


# ==================== ЗАПУСК ПРОГРАММЫ ====================
def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    # Устанавливаем иконку приложения
    if os.path.exists("logo.png"):
        app.setWindowIcon(QIcon("logo.png"))
    else:
        app.setWindowIcon(QIcon.fromTheme("office-chart-line"))
    window = MainWindow()
    window.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()